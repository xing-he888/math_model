from langgraph.graph import StateGraph,START,END
from langgraph.graph.message import add_messages
from typing import TypedDict,List,Dict,Annotated
from langgraph.graph import MessagesState
from langgraph.types import interrupt, Command
from loguru import logger
from pathlib import Path
import sys
import os

# 无论从根目录还是 src 目录运行，都确保能找到 tool.py
sys.path.insert(0, str(Path(__file__).resolve().parent))

from langchain_core.messages import HumanMessage,AIMessage, SystemMessage, ToolMessage, RemoveMessage
from langchain_core.callbacks import BaseCallbackHandler
from dotenv import  load_dotenv
from langgraph.prebuilt import ToolNode
from langchain_core.tools import tool
from models import get_model, get_struct_model, list_models, MODEL_REGISTRY, DEFAULT_MODEL
from pypdf import PdfReader
from langgraph.checkpoint.memory import InMemorySaver
from langgraph.types import Send
from tool import tools as base_tools, write_workspace, WORKSPACE_ROOT, set_workspace, get_workspace, ws_root
from src.skills import load_skill_tools
from src.skills.writing import read_writing_skill

# 基础工具(tool.py)+ 可插拔 skill(src/skills/):新增 skill 只需在 src/skills/ 放一个模块并导出 tools 列表
tools = base_tools + load_skill_tools()

TOOLS_BY_NAME = {t.name: t for t in tools}
import pandas as pd
import json
import re
import time
import threading
import operator
import subprocess
import shutil

# 然后将 search 作为工具传递给 LangGraph 的 Agent
#读取env文件中的apikey
load_dotenv(override=True)

# ---------- 模型（可配置：从注册表选择，支持 deepseek/gpt/glm/qwen/kimi/mimo） ----------
# 通过环境变量 MATH_MODEL 选择要用的模型，例如 deepseek / gpt / glm / qwen / kimi / mimo
# 新增模型：在 src/models.py 的 MODEL_REGISTRY 里加一项即可，无需改这里。
# ---------- 双实例：干活(不思考) vs 动脑(可思考) ----------
# 背景: DeepSeek V4 默认思考模式, 且思考模式下拒绝 tool_choice="required"/指定函数名(HTTP 400),
# 因此工具循环必须用不思考实例; 其余节点(建模分析、诊断、论文写作等)默认走可思考实例。
# 注意: 结构化输出(格式化)步骤不在此列——它们固定走 DeepSeek 非思考专用实例(见 struct_model),
# 不随 set_model 切换, 因为强制思考模型会拒绝 with_structured_output 的强制 tool_choice。
model_tool = get_model(role="tool")
model_text = get_model(role="text")
model = model_tool                      # 旧变量名保留, 指向非思考实例(worker 兜底/标题生成等)
model_with_tool=model_tool.bind_tools(tools=tools)
model_think_tool=model_text.bind_tools(tools=tools)
tool_node=ToolNode(tools=tools)

# 当前生效模型的显示名(按角色),供日志打印"本次调用用了哪个模型";
# tool/text 随 set_model 更新,struct(格式化)固定 deepseek 不随切换变化
MODEL_NAMES = {"tool": DEFAULT_MODEL, "text": DEFAULT_MODEL, "struct": "deepseek"}

def _llm_role(m) -> str:
    """按实例身份推断角色: text(思考)/struct(格式化固定)/其余归 tool(干活)。
    注意 with_structured_output/bind_tools 都会生成新包装对象,须逐一身份比对。"""
    if m is model_text or m is model_think_tool:
        return "text"
    if m is struct_model or m is model_with_struct or m is model_feedback_struct or m is model_plan_struct:
        return "struct"
    return "tool"

def _log_llm(m, action: str = "") -> None:
    role = _llm_role(m)
    suffix = f" · {action}" if action else ""
    logger.info(f"LLM调用 → {MODEL_NAMES.get(role, '?')}[{role}]{suffix}")

# ---------- Token 消耗 / 缓存命中统计(后端 GET /api/usage 读取, 前端侧边栏展示) ----------
# 全局单桶:与"手动单题运行"用法一致, /api/reset 与每次新运行时由后端清零。
# 节点由 langgraph 放线程池并行执行, 累加必须持锁; 统计环节任何异常都吞掉, 绝不影响主流程。
_usage_lock = threading.Lock()
usage_stats = {
    "calls": 0,
    "prompt_tokens": 0,
    "completion_tokens": 0,
    "cache_hit_tokens": 0,    # DeepSeek: prompt_cache_hit_tokens; OpenAI: prompt_tokens_details.cached_tokens
    "cache_miss_tokens": 0,   # DeepSeek 专有
    "cache_supported": False, # 本次累计中是否出现过缓存字段(决定前端显示命中率还是"—")
    "by_role": {},            # 角色(tool/text/struct)分桶: {calls, prompt_tokens, completion_tokens}
}

def _record_usage(msg, role: str) -> None:
    """从底层 AIMessage 的 token_usage 累计一笔。缓存命中字段按优先级探测:
    DeepSeek(prompt_cache_hit_tokens) → OpenAI(prompt_tokens_details.cached_tokens);
    都没有的模型命中数保持 0 且不置 cache_supported, 前端显示"—", 厂商未来上报后自动点亮。"""
    try:
        usage = (getattr(msg, "response_metadata", None) or {}).get("token_usage") or {}
        um = getattr(msg, "usage_metadata", None) or {}
        prompt = int(usage.get("prompt_tokens") or um.get("input_tokens") or 0)
        completion = int(usage.get("completion_tokens") or um.get("output_tokens") or 0)
        if "prompt_cache_hit_tokens" in usage:                 # DeepSeek
            supported, hit = True, int(usage.get("prompt_cache_hit_tokens") or 0)
            miss = int(usage.get("prompt_cache_miss_tokens") or 0)
        elif isinstance(usage.get("prompt_tokens_details"), dict):  # OpenAI 系
            supported, hit = True, int(usage["prompt_tokens_details"].get("cached_tokens") or 0)
            miss = 0
        else:
            supported, hit, miss = False, 0, 0
        if prompt <= 0 and completion <= 0:
            return
        with _usage_lock:
            s = usage_stats
            s["calls"] += 1
            s["prompt_tokens"] += prompt
            s["completion_tokens"] += completion
            s["cache_hit_tokens"] += hit
            s["cache_miss_tokens"] += miss
            s["cache_supported"] = s["cache_supported"] or supported
            r = s["by_role"].setdefault(role, {"calls": 0, "prompt_tokens": 0, "completion_tokens": 0})
            r["calls"] += 1
            r["prompt_tokens"] += prompt
            r["completion_tokens"] += completion
    except Exception:
        pass

class _UsageCollector(BaseCallbackHandler):
    """单次 invoke 的用量采集回调: 在 on_llm_end 抓底层 AIMessage 记账。
    with_structured_output 的返回值是解析后的 dict(元数据被吃掉), 必须经回调拿原始消息;
    普通调用统一也走这里, 保证全项目记账口径只有一条。"""
    def __init__(self) -> None:
        super().__init__()
        self.msgs = []

    def on_llm_end(self, response, **kwargs) -> None:
        try:
            for gens in (getattr(response, "generations", None) or []):
                for g in (gens or []):
                    m = getattr(g, "message", None)
                    if m is not None:
                        self.msgs.append(m)
        except Exception:
            pass

    def drain(self, role: str) -> None:
        for m in self.msgs:
            _record_usage(m, role)
        self.msgs = []

def usage_snapshot() -> dict:
    """返回统计副本(嵌套 by_role 一并复制), 供后端接口安全序列化"""
    with _usage_lock:
        return {**usage_stats, "by_role": {k: dict(v) for k, v in usage_stats["by_role"].items()}}

def reset_usage() -> None:
    """清零统计: 后端 /api/reset 与每次新运行时调用"""
    with _usage_lock:
        usage_stats.update({
            "calls": 0, "prompt_tokens": 0, "completion_tokens": 0,
            "cache_hit_tokens": 0, "cache_miss_tokens": 0, "cache_supported": False,
        })
        usage_stats["by_role"] = {}

# ---------- AI 使用事件日志(2026 国赛 AI 使用规定: 声明与《AI 工具使用详情》的事实来源) ----------
# 与 usage_stats 同款模式: 模块级+锁, 新运行时随 reset_usage 一并清零; 记录失败不影响主流程。
_ai_log_lock = threading.Lock()
ai_usage_events: List[Dict[str, str]] = []   # {time, kind, model, summary, detail}

def log_ai_event(kind: str, summary: str, model: str = "", detail: str = "") -> None:
    """记录一条 AI 使用/人工审查事件。kind: '调用' | '人工审查'; 超过 500 条后丢弃(防撑爆内存)。"""
    try:
        entry = {
            "time": time.strftime("%Y-%m-%d %H:%M:%S"),
            "kind": kind,
            "model": model or "",
            "summary": (summary or "")[:120],
            "detail": (detail or "")[:400],
        }
        with _ai_log_lock:
            if len(ai_usage_events) < 500:
                ai_usage_events.append(entry)
    except Exception:
        pass

def ai_events_snapshot() -> List[Dict[str, str]]:
    with _ai_log_lock:
        return [dict(e) for e in ai_usage_events]

def reset_ai_log() -> None:
    with _ai_log_lock:
        ai_usage_events.clear()

def _first_human_excerpt(msgs) -> str:
    """取第一条用户消息做摘录(AI 详情 PDF 里'典型提示摘录'的素材)。"""
    try:
        for m in (msgs or []):
            if getattr(m, "type", "") == "human":
                return re.sub(r"\s+", " ", str(getattr(m, "content", "")))[:400]
    except Exception:
        pass
    return ""

def _invoke_llm(model_obj, msgs, action: str = "", retries: int = 2):
    """模型调用带降级 + 重试：
    ① 思考实例被 API 以 400/tool_choice/thinking 拒绝 → 确定性错误，立即降级为
       非思考实例（降级实例递归复用本函数，享受同样的日志与重试；因身份检查只对
       思考实例成立，递归深度最多 1 层，必然终止）；
    ② 401/403 鉴权错误 → 重试无意义，快速抛出；
    ③ 其余瞬时错误（429 限流/超时/5xx/连接重置等）→ 退避 2s/4s 重试 retries 次，
       仍失败原样抛出。
    注意用 id 严格判断实例身份——建模节点超限路径传入的 tool_choice="none" 实例
    绝不允许被降级替换，否则会重新放开工具调用、破坏防死循环保证。
    sync 节点由 langgraph 放在线程池 executor 执行，重试里的 sleep 只阻塞本分支
    线程，不影响事件循环、SSE 流与其它并行分支。"""
    _log_llm(model_obj, action)
    for attempt in range(retries + 1):
        try:
            cb = _UsageCollector()
            resp = model_obj.invoke(msgs, config={"callbacks": [cb]})
            cb.drain(_llm_role(model_obj))   # token/缓存记账, 统计异常已在内部吞掉
            log_ai_event("调用", action or "通用调用",
                         MODEL_NAMES.get(_llm_role(model_obj), ""), _first_human_excerpt(msgs))
            return resp
        except Exception as e:
            msg = str(e)
            # ① 思考模式被拒：确定性错误，直接降级换实例
            if any(model_obj is m for m in (model_text, model_think_tool)) and (
                "400" in msg or "tool_choice" in msg.lower() or "thinking" in msg.lower()
            ):
                logger.warning(f"思考模式调用被拒({e})，自动降级为非思考实例重试")
                return _invoke_llm(model_with_tool, msgs, action, retries)
            # ② 鉴权错误：重试无意义
            if "401" in msg or "403" in msg:
                raise
            # ③ 瞬时故障：退避重试
            if attempt < retries:
                logger.warning(f"LLM 瞬时故障({e})，{2 * (attempt + 1)}s 后重试({attempt + 2}/{retries + 1})")
                time.sleep(2 * (attempt + 1))
            else:
                raise

# 运行时动态切换全局模型（前端/API 可调用），无需重启进程。
# 只重设思考/干活两套实例及其派生; struct_model(格式化专用)固定 DeepSeek, 刻意不重建。
def set_model(name: str = None) -> str:
    """重设全部可切换的模型全局对象，返回实际生效的 key。"""
    global model, model_tool, model_text, model_with_tool, model_think_tool
    key = (name or os.getenv("MATH_MODEL", DEFAULT_MODEL)).strip().lower()
    if key not in MODEL_REGISTRY:
        key = DEFAULT_MODEL
    model_tool = get_model(key, role="tool")
    model_text = get_model(key, role="text")
    model = model_tool
    model_with_tool = model_tool.bind_tools(tools=tools)
    model_think_tool = model_text.bind_tools(tools=tools)
    MODEL_NAMES["tool"] = MODEL_NAMES["text"] = key
    return key

# 题目输入输出目录：随当前题目工作区（workspaces/{题目id}/）动态解析，多题互不干扰
def question_dir() -> Path:
    return ws_root() / "question"

def dataset_dir() -> Path:
    return ws_root() / "dataset"

# redo 清场哨兵:operator.add 下空列表清不掉旧值(旧值+[]=旧值),必须用自定义 reducer 识别 RESET 强制归零
# 注意:langgraph 1.2.x 的 InMemorySaver 会把 channel 写入 msgpack 序列化后存检查点,
# object() 哨兵不可序列化会直接 TypeError;必须用可序列化的字符串哨兵。
# 且 resume 重放时写入值经"序列化→反序列化"往返,回来的是值相等的新对象,
# 因此 reducer 里只能用 == 比较而不能用 is。
RESET = "__RESET__"

def add_or_reset(left: list, right) -> list:
    """可重置的追加 reducer:收到 RESET 哨兵时清空,否则按 operator.add 追加"""
    if right == RESET:
        return []
    return left + (right or [])

#全局状态
class over_all_state(MessagesState):
    input_problem: Annotated[str, '题目原始文本']
    problem_str: Annotated[str, '问题题干']
    problem_index: Annotated[Dict[str, str], '问题索引字典']
    dataset: Annotated[str, '题目给出的数据集']
    modeling_approach: Annotated[str, '建模手自己的思路']
    modeling_analysis: Annotated[str, '模型分析出的基础思路']
    plan_struct: Annotated[str, '结构化建模方案(JSON文本:problem_type/objective/constraints等)'] = ''
    # methods 用覆盖语义(无 reducer):modeling 每次全量重算候选方法集,打回重做时不会累积旧值
    # 形态:modeling 产出初版(List[str] 查表兜底),debate_plan 升级为方法卡列表
    # (List[dict]: name/paradigm/rationale/assumption/tools);下游节点兼容两种形态。
    methods: List[str] = []
    # 建模辩论未决分歧(issue/adopter/status/note),随 Send 下发各方法作为可发挥空间
    disagreements: Annotated[List[dict], '建模辩论未决分歧,随 Send 下发'] = []
    retry_count: Annotated[int, '工具重试计数'] = 0
    tool_rounds: Annotated[int, '工具调用总轮数'] = 0
    review_feedback: Annotated[str, '审核意见'] = ''
    review_result: Annotated[str, '审核结果'] = ''
    human_feedback: Annotated[str, '人工对最终总结的意见'] = ''
    method: Annotated[str, 'worker方法身份'] = ''
    dataset_files: Annotated[str, '数据集文件名清单'] = ''
    answers: Annotated[List[Dict[str, str]], add_or_reset] = []
    code_files: Annotated[List[str], add_or_reset] = []
    run_report: List[Dict[str, str]] = []
    failed_qs: Annotated[List[str], add_or_reset] = []
    done_pairs: Annotated[List[str], add_or_reset] = []
    # 注意:必须用 add_messages 而非 operator.add——final_analysis 打回时会写入 RemoveMessage 删除指令,
    # 只有 add_messages 能消化它;operator.add 会把指令本身拼进列表,后续 invoke 模型时抛 TypeError
    compare_msgs: Annotated[List, add_messages] = []
    final_summary: Annotated[str, '最终总结'] = ''
    article_chapters: Annotated[List[str], '已生成的论文章节'] = []
    compile_status: Annotated[str, '论文编译验证状态'] = ''
    model_iteration: Annotated[int, '建模迭代次数(feedback_check 打回计数)'] = 0
    feedback_notes: Annotated[List[str], add_or_reset] = []  # 上轮质检失败原因,modeling 节点据此修正思路

#定义格式化的状态，仅开始时使用
class structed_output_state(TypedDict):
    problem_str:Annotated[str,'问题题干']
    problem_index:Annotated[Dict[str, str], "问题索引字典"]

# ---------- 格式化专用实例：固定 DeepSeek 非思考，不随前端切换模型变化 ----------
# with_structured_output 内部会发强制 tool_choice，强制思考模型(如 OpenRouter 接入的
# reasoning 模型)会 HTTP 400 拒绝; 问题提取/质检这类格式化步骤钉死 DeepSeek,
# 前端切换的模型只影响其余思考型调用。需在 .env 配置 DEEPSEEK_API_KEY(前端已注明)。
struct_model = get_struct_model()
model_with_struct = struct_model.with_structured_output(schema=structed_output_state)

# feedback_check 质检节点的裁决结构(同样走固定格式化实例)
class FeedbackVerdict(TypedDict):
    passed: Annotated[bool, '建模方案与求解结果是否通过质检']
    reason: Annotated[str, '不通过时的具体原因']
    suggestion: Annotated[str, '打回重做时给建模手的修正建议']

model_feedback_struct = struct_model.with_structured_output(schema=FeedbackVerdict)

# debate_plan 收口:把修订后的建模方案强制结构化(格式硬保证,不走围栏解析软路)
class PlanStruct(TypedDict):
    problem_type: Annotated[str, '题目类型(优化/预测/评价/分类/机理,单选)']
    variables: Annotated[List[str], '决策变量及含义']
    objective: Annotated[str, '目标函数表达式或文字描述']
    constraints: Annotated[List[str], '约束条件列表']
    per_question_method: Annotated[Dict[str, str], '每问采用的方法']

model_plan_struct = struct_model.with_structured_output(schema=PlanStruct)

#质检节点:求解结果汇总后、汇合前,用结构化输出判断建模方案是否可信
#不通过则打回 modeling 重做(清空 operator.add 系字段防止污染),最多迭代 3 次强制放行
def feedback_check(state: over_all_state):
    logger.info('正在运行 feedback_check 节点')
    iteration = state.get("model_iteration", 0)
    answers = state.get("answers") or []

    # 无解或迭代达上限:直接放行,防死循环
    if not answers or iteration >= 3:
        if iteration >= 3:
            logger.warning('建模迭代已达 3 次上限，强制放行进入后续流程')
        # 放行时同步清残留,避免旧质检状态污染后续轮次
        return Command(goto="collect_branches", update={"feedback_notes": RESET, "model_iteration": 0})

    problem_str = state["problem_str"]
    analysis = state.get("modeling_analysis") or ""
    # answers 结构是 [{方法名: 结果文本}, ...]——先按方法名合并,再取文本摘要;
    # 旧实现读 a['question']/a['answer'] 与真实结构不符,质检摘要恒为空,已修正
    merged = {}
    for item in answers:
        for k, v in item.items():
            merged.setdefault(k, []).append(v)
    brief = "\n".join(
        f"- 方法《{k}》: {str(text)[:500]}"
        for k, texts in merged.items() for text in texts[:3]
    ) or "(无结果摘要)"
    failed = state.get("failed_qs") or []
    failed_str = "；".join(failed) if failed else "无"

    prompt = (
        "你是数学建模质检员。下面是建模方案与各小问的求解结果摘要，请判断方案与结果是否可信、自洽。\n"
        f"【题目】{problem_str}\n"
        f"【建模思路】{analysis}\n"
        f"【失败的小问】{failed_str}\n"
        f"【求解结果摘要】\n{brief}\n\n"
        "判定标准：结果明显矛盾、大量小问失败、模型与题意不符才算不通过；结论合理即可通过。"
    )
    try:
        verdict = _invoke_llm(model_feedback_struct, prompt, action="建模质检")
    except Exception as e:
        logger.warning(f'feedback_check 质检调用失败，默认放行: {e}')
        return Command(goto="collect_branches", update={"feedback_notes": RESET, "model_iteration": 0})

    passed = bool(verdict.get("passed", True))
    reason = str(verdict.get("reason", ""))
    suggestion = str(verdict.get("suggestion", ""))

    if passed:
        logger.info('feedback_check 质检通过')
        # 通过即归零:迭代计数与失败原因不带入下一题/下一轮
        return Command(goto="collect_branches", update={"feedback_notes": RESET, "model_iteration": 0})

    logger.warning(f'feedback_check 质检不通过，打回 modeling 重做(第 {iteration + 1} 次): {reason}')
    return Command(
        goto="modeling",
        update={
            "answers": RESET,
            "code_files": RESET,
            "failed_qs": RESET,
            "done_pairs": RESET,
            "disagreements": [],   # 覆盖语义无 reducer,清空只能用 [];防上轮辩论的旧分歧点残留给本轮 worker
            "review_feedback": "",       # 清掉旧审核意见,避免 modeling 误判为审核打回
            "retry_count": 0,
            "tool_rounds": 0,
            "model_iteration": iteration + 1,
            "feedback_notes": [f"第{iteration + 1}次质检不通过：{reason} 修正建议：{suggestion}"],
        },
    )


#用来初始化input_problem
def load_problem(state: over_all_state) -> dict:
    """读取 question 文件夹里的题目文件（.txt/.md 按文本读，.pdf 用 pypdf 提取），合并后写入 input_problem 字段"""
    logger.info('正在运行load_problem节点')
    path = question_dir()
    if not path.is_dir():
        return {"input_problem": "(question 文件夹为空)"}
    contents = []
    for item in path.iterdir():
        if not item.is_file():
            continue
        try:
            if item.suffix.lower() == ".pdf":
                reader = PdfReader(str(item))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            else:
                text = item.read_text(encoding="utf-8")
        except Exception as e:
            text = f"(读取 {item.name} 失败: {e})"
        contents.append(f"=== {item.name} ===\n{text}")
    return {"input_problem": "\n\n".join(contents) if contents else "(question 文件夹为空)"}

#从 LLM 文本输出中提取 JSON 对象(容忍 ```json 围栏与前后杂文),失败抛 ValueError
def _extract_json_obj(text: str) -> dict:
    s = (text or "").strip()
    s = re.sub(r"^```(?:json)?\s*", "", s)
    s = re.sub(r"\s*```$", "", s)
    l, r = s.find("{"), s.rfind("}")
    if l == -1 or r <= l:
        raise ValueError("输出中未找到 JSON 对象")
    obj = json.loads(s[l:r + 1])
    if not isinstance(obj, dict):
        raise ValueError("JSON 不是对象")
    return obj

def _extract_last_json_fence(text: str) -> tuple[dict, str]:
    """从文本的 ``` 围栏块中自后向前找第一个能整体解析成 JSON 对象的块,
    返回 (方案, 剥掉该块后的正文);找不到抛 ValueError。
    不锚定输出末尾:模型在围栏后补总结语、正文含 LaTeX 花括号都不影响;
    自后向前也保证正文更早处的 JSON 样例块不会被误当成方案。"""
    parts = re.split(r"(```[a-zA-Z0-9]*\s*)", text or "")
    # parts 偶数位是围栏外正文,奇数位是围栏标记(开栏/闭栏交替,开栏在奇数位)
    for j in range(len(parts) - 2, 0, -1):
        if j % 2 == 0:
            continue
        content = parts[j + 1].strip() if j + 1 < len(parts) else ""
        if not (content.startswith("{") and content.endswith("}")):
            continue
        try:
            obj = json.loads(content)
        except Exception:
            continue
        if not isinstance(obj, dict):
            continue
        # 整块 = 开栏标记(j) + 块内容(j+1) + 闭栏标记(j+2);未闭合的尾块同样剥除
        cleaned = "".join(parts[:j]) + "".join(parts[j + 3:])
        return obj, cleaned
    raise ValueError("围栏块中未找到可解析的 JSON 对象")

#对文本进行格式化用来初始化problem_str和problem_index。
#固定走 struct_model(DeepSeek 非思考)——不随前端切换模型变化,规避强制思考模型
#拒绝 with_structured_output 强制 tool_choice 的 400;三层降级保证流程不中断
def question_structed(state:over_all_state) ->structed_output_state:
    logger.info('正在运行question_structed节点')
    input_problem=state['input_problem']
    prompt = (
        "数学建模题目，提取出题目中所有编号问题的完整题干"
        "（从“问题 X：”开始到下一个“问题”、“问题 X”或“相关说明”之前），"
        "问题数量不固定，以题目实际编号为准，如“问题1”、“问题2”……\n"
        "只输出一个 JSON 对象（不要代码围栏与解释文字），格式：\n"
        '{"problem_str": "完整题目全文", "problem_index": {"问题1": "该问完整题干", "问题2": "..."}}\n\n'
        f"题目内容：\n{input_problem}"
    )

    def _finalize(p_str: str, p_index, note: str) -> dict:
        # 索引必须非空兜底:空字典会让下游视为"无小问",静默跳过四方法求解
        index = p_index or {}
        if not index and p_str:
            index = {"问题1": p_str}
            note += "；problem_index 为空，已把整题作为「问题1」兜底"
        return {"problem_str": p_str or input_problem,
                "problem_index": index,
                "messages": [AIMessage(content=note)]}

    # 第一层:结构化输出(固定 DeepSeek 非思考实例)
    try:
        _log_llm(model_with_struct, "问题提取·结构化")
        cb = _UsageCollector()
        resp = model_with_struct.invoke([HumanMessage(prompt)], config={"callbacks": [cb]})
        cb.drain("struct")
        log_ai_event("调用", "问题提取·结构化", "deepseek", _first_human_excerpt([HumanMessage(prompt)]))
        if isinstance(resp, dict) and resp.get("problem_str"):
            return _finalize(str(resp["problem_str"]), resp.get("problem_index"), "已成功提取问题索引")
        logger.warning(f"question_structed 结构化返回异常(resp={resp!r})，转手动解析")
    except Exception as e:
        logger.warning(f"question_structed 结构化调用失败({e})，回退普通文本调用")

    # 第二层:同一 DeepSeek 实例的普通文本调用 + 手动 JSON 解析
    try:
        _log_llm(struct_model, "问题提取·文本回退")
        cb = _UsageCollector()
        raw_resp = struct_model.invoke([HumanMessage(prompt)], config={"callbacks": [cb]})
        cb.drain("struct")
        log_ai_event("调用", "问题提取·文本回退", "deepseek", _first_human_excerpt([HumanMessage(prompt)]))
        raw = raw_resp.content
        data = _extract_json_obj(raw)
        if data.get("problem_str"):
            return _finalize(str(data["problem_str"]), data.get("problem_index"), "已通过文本模式提取问题索引")
        logger.warning(f"question_structed 手动解析缺少 problem_str(data={data!r})")
    except Exception as e:
        logger.warning(f"question_structed 文本回退解析失败({e})，走整题单问兜底")

    # 第三层:整题单问兜底,保证求解全流程不被跳过
    return _finalize(input_problem, {}, "问题索引提取失败，已把整题作为单问回退")

#建模路由:工具调用未超限才进工具节点;超限后即使模型仍带 tool_calls 也强制进入审核,防止死循环
def modeling_route(state: over_all_state) -> str:
    last = state["messages"][-1] if state["messages"] else None
    if getattr(last, "tool_calls", None) and state.get("tool_rounds", 0) < MAX_TOOL_ROUNDS:
        return "tools"
    return "review"

#此节点为中断节点建模手提供对题目的看法以及让模型给出自己的思路
def modeling(state: over_all_state) -> over_all_state:
    logger.info('正在运行 modeling 节点')
    
    # ----- 1. 重试逻辑处理 -----
    last_message = state["messages"][-1] if state["messages"] else None
    retry_count = state.get("retry_count", 0)
    tool_rounds = state.get("tool_rounds", 0)
    is_tool_return = isinstance(last_message, ToolMessage)
    is_rework = bool(state.get("review_feedback"))

    # 工具每返回一次(无论成败)都计一轮,防止模型反复成功调用工具导致死循环
    if is_tool_return:
        tool_rounds += 1
        content = last_message.content
        if "错误" in content or "失败" in content or "不存在" in content:
            retry_count += 1
            logger.warning(f"工具执行失败，正在进行第 {retry_count} 次重试...")

    # 失败重试超过 3 次或工具总轮数超过上限时,强制禁用工具
    if retry_count >= 3 or tool_rounds >= MAX_TOOL_ROUNDS:
        logger.warning(f"工具调用达上限(失败{retry_count}次/共{tool_rounds}轮)，将禁用工具调用，强制输出结果")
        model_to_use = model_with_tool.bind_tools(tools, tool_choice="none")
        retry_hint = "\n（注意：工具调用次数已达上限，请不要再调用工具，直接根据已有信息给出最终分析。）"
    else:
        model_to_use = model_think_tool
        retry_hint = f"\n（当前工具已调用 {tool_rounds} 轮，失败 {retry_count} 次；如无需更多信息请直接给出最终分析。）"

    # 首次进入节点时中断询问建模手思路；工具调用回退、审核打回或质检打回时不再中断，直接继续分析
    if is_tool_return or is_rework or bool(state.get("feedback_notes")):
        modeling_approach = ""
    else:
        modeling_approach = interrupt('请简述你的建模思路（可直接回车跳过）') or ""
        log_ai_event("人工审查", f"建模思路人工确认: {'提供了人工思路' if modeling_approach.strip() else '跳过,由模型自行构思'}")
    problem_str = state["problem_str"]
    problem_index = state["problem_index"]
    dataset = state.get("dataset")

    system_prompt = (
        "你是一名顶级的数学建模手。请严格遵循以下步骤对题目进行深度分析：\n"
        "1. 问题定性（优化/预测/评价/分类/机理）\n"
        "2. 数据洞察与预处理建议\n"
        "3. 合理假设\n"
        "4. 数学模型构建（目标函数+约束条件）\n"
        "5. 求解算法推荐\n"
        "6. 优劣势分析\n"
        "题目与数据已完整提供，通常无需调用工具；如需查资料请用 search。\n"
        "若调用 read_workspace，work_dir 参数必填，只能是 code/paper/photo/dataset 之一，例如 read_workspace(work_dir=\"paper\", rel_path=\"问题1_思路.md\")。\n"
        "若工具连续报错，请根据你的常识完成分析。\n"
        "【重要·结构化方案】在分析文字的最后，用 ```json 代码围栏额外输出一份结构化建模方案"
        "（供后续求解节点直接消费，务必与你的分析一致），格式：\n"
        "{\n"
        '  "problem_type": "题目类型(优化/预测/评价/分类/机理，单选)",\n'
        '  "variables": ["决策变量1及含义", "决策变量2及含义"],\n'
        '  "objective": "目标函数表达式(如 max 3x1+4x2)",\n'
        '  "constraints": ["约束1表达式(如 2x1+4x2<=200)", "约束2表达式"],\n'
        '  "per_question_method": {"问题1": "该问建议采用的方法", "问题2": "..."}\n'
        "}\n"
        "若无法给出具体表达式，objective/constraints 填简要文字描述即可，不得省略整个 JSON。\n"
        "输出完上述 ```json 围栏后立即结束本次回复，围栏之后不要再输出任何文字。"
    )

    user_parts = [
        f"每一个小问的题目：{problem_index}",
        f"题目内容：{problem_str}",
        f"数据集位置：{dataset}" if dataset else "（无数据集）"
    ]
    if is_tool_return:
        user_parts.append(f"最近一次工具调用的返回结果（请据此判断上一步是否成功、如何修正参数）：\n{last_message.content}")
    elif modeling_approach and modeling_approach.strip():
        user_parts.append(f"用户提供的思路：{modeling_approach}")
    elif is_rework:
        user_parts.append(f"建模手审核意见（请据此重新分析）：{state.get('review_feedback')}")
    else:
        user_parts.append("（用户未提供思路，请自行构思）")
    # 质检打回:附上前轮失败原因,要求建模手针对性修正方案
    feedback_notes = state.get("feedback_notes") or []
    if feedback_notes:
        user_parts.append(
            "【重要】此前建模方案未通过质检被打回，历史失败原因如下（请务必针对性修正，不要重复同样的思路）：\n"
            + "\n".join(f"- {note}" for note in feedback_notes)
        )
    user_parts.append(retry_hint)
    user_message = "\n".join(user_parts)

    response = _invoke_llm(model_to_use, [
        SystemMessage(content=system_prompt),
        HumanMessage(content=user_message)
    ], action="建模分析")

    # 提取结构化方案:先扫 ``` 围栏自后向前取方案(容错围栏后补话/正文花括号),
    # 失败再退回通用提取;成功时把方案块从正文剥掉,JSON 只经 plan_struct 下发
    plan_struct, methods = "", []
    analysis_text = response.content
    try:
        try:
            plan, analysis_text = _extract_last_json_fence(response.content)
        except ValueError:
            plan = _extract_json_obj(response.content)
        ptype = str(plan.get("problem_type") or "").strip()
        methods = TYPE_METHODS.get(ptype) or METHODS_FALLBACK
        plan_struct = json.dumps(plan, ensure_ascii=False, indent=2)
        logger.info(f"modeling 结构化为类型「{ptype}」, 派发方法: {methods}")
    except Exception as e:
        logger.warning(f"modeling 结构化方案解析失败({e})，将沿用固定方法集")

    return {
        "messages": [response],          # 追加 AI 回复
        "modeling_analysis": analysis_text,  # 已剥掉方案围栏,JSON 只走 plan_struct
        "plan_struct": plan_struct,      # 结构化方案(JSON文本), 供 worker/论文直接消费
        "methods": methods or METHODS_FALLBACK,  # 本轮的候选方法集(类型感知)
        "retry_count": retry_count,      # 更新重试计数器
        "tool_rounds": tool_rounds       # 更新工具轮数
    }


#此函数用于转换格式
def _read_file_by_suffix(path: Path) -> str:
    """按文件后缀分发读取：表格类（csv/tsv/xlsx）转 DataFrame 文本、pdf 提取、docx 提取、
    json/jsonl 美化输出，其余按 utf-8 文本读取"""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix in (".csv", ".tsv"):
        df = pd.read_csv(path, sep="\t" if suffix == ".tsv" else ",", encoding="utf-8")
        return df.fillna("").to_string(index=False)
    if suffix in (".xlsx", ".xls"):
        sheets = []
        for sheet_name in pd.ExcelFile(path).sheet_names:
            df = pd.read_excel(path, sheet_name=sheet_name)
            sheets.append(f"### 工作表: {sheet_name} ###\n" + df.fillna("").to_string(index=False))
        return "\n\n".join(sheets)
    if suffix == ".json":
        return json.dumps(json.loads(path.read_text(encoding="utf-8")), ensure_ascii=False, indent=2)
    if suffix == ".jsonl":
        lines = [line for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
        return "\n".join(json.dumps(json.loads(line), ensure_ascii=False) for line in lines)
    if suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            parts.append(pd.DataFrame(rows).fillna("").to_string(index=False))
        return "\n".join(parts)
    if suffix in (".txt", ".md", ".xml", ".html", ".log", ".dat", ".cfg", ".ini"):
        return path.read_text(encoding="utf-8")
    # 未知格式回退 utf-8 文本读取
    return path.read_text(encoding="utf-8")

#此节点用来读取dataset文件夹中文件的数据
def read_dataset(state:over_all_state) -> over_all_state:
    """读取 dataset 文件夹里的数据文件（按后缀分发处理多种格式），合并后写入 dataset 字段"""
    logger.info('正在运行read_dataset节点')
    path = dataset_dir()
    if not path.is_dir():
        return {"dataset": "(dataset 文件夹为空)"}
    contents = []
    for item in sorted(path.iterdir()):
        if not item.is_file():
            continue
        try:
            text = _read_file_by_suffix(item)
        except UnicodeDecodeError:
            text = f"(文件 {item.name} 是二进制格式（{item.suffix}），无法直接读取，请先转为 csv/txt 放入 dataset 文件夹)"
        except Exception as e:
            text = f"(读取 {item.name} 失败: {e})"
        contents.append(f"=== {item.name} ===\n{text}")
    return {"dataset": "\n\n".join(contents) if contents else "(dataset 文件夹为空)"}


#此节点用于人工展示与审核 modeling_analysis：可通过、打回重新建模、或带建议重新建模
def review_modeling_analysis(state: over_all_state) -> Command:
    logger.info('正在运行 review_modeling_analysis 节点')
    analysis = state.get("modeling_analysis") or "(暂无建模分析结果)"
    feedback = interrupt(
        f"建模手请审核以下建模分析结果：\n\n{analysis}\n\n"
        "审核意见（直接回车=通过；输入\"打回\"=不通过重新分析；输入其他内容=作为建议重新分析）："
    )
    feedback = (feedback or "").strip()
    if not feedback or feedback in ("通过", "同意", "ok", "OK", "好", "可以"):
        logger.info('建模手审核通过，进入下一环节')
        log_ai_event("人工审查", "建模分析人工审核: 通过")
        return Command(update={"review_feedback": "", "review_result": "passed", "retry_count": 0, "tool_rounds": 0}, goto="send_problem_index")
    logger.info(f"建模手打回，审核意见：{feedback}")
    log_ai_event("人工审查", f"建模分析人工审核: 打回重做, 意见: {feedback[:120]}")
    # disagreements 随被丢弃的旧方案一并清空:它只在 debate_plan 写入,若重跑时辩论
    # 静默降级,残留的旧分歧点会配着新方案下发给 worker(与质检打回同一处理)
    return Command(update={"review_feedback": feedback, "review_result": "rework", "retry_count": 0, "tool_rounds": 0, "disagreements": []}, goto="modeling")


# 候选方法集:按题目类型(problem_type)选派,替代"固定四法硬跑",避免方法凑数。
# METHODS_FALLBACK 兜底:类型判定缺失/失败时沿用原固定四法,保证流程不中断。
TYPE_METHODS = {
    "优化": ["精确规划", "启发式优化", "数值枚举"],
    "预测": ["机理建模", "统计回归", "机器学习"],
    "评价": ["层次分析", "熵值法", "模糊综合评价"],
    "分类": ["判别分类", "聚类分析"],
    "机理": ["微分方程解析", "数值仿真"],
}
# 兜底方法集:类型判定失败时启用,选三个"几乎万能"的方法,保证任何题都有可用求解路径
METHODS_FALLBACK = ["机理建模", "统计建模", "智能优化"]

#建模阶段工具调用总轮数上限,防止模型反复调工具死循环
MAX_TOOL_ROUNDS = 6

# ---------- 建模辩论节点：质疑者攻击初稿 → 修订者收敛 → 收口保格式 ----------
# 目标:给"决策层"加独立于生成者的审阅——质疑者独立读题面攻击分类/假设/目标函数,
# 修订者逐条回应产出修订版 plan_struct + 方法卡集 + 分歧点。
# 铁律:任一步失败(调用/解析/校验)都静默降级——原样返回 modeling 初稿,行为等同改前,
# 绝不让新节点成为全链路的新单点。
def debate_plan(state: over_all_state) -> dict:
    logger.info('正在运行 debate_plan 节点')
    draft = state.get("plan_struct") or ""
    if not draft:
        logger.warning('debate_plan: 无建模初稿,直接放行')
        return {}
    problem_str = state.get("problem_str") or ""
    problem_index = state.get("problem_index") or {}
    analysis = state.get("modeling_analysis") or ""
    initial_methods = state.get("methods") or METHODS_FALLBACK

    # ---- 1. 质疑者:独立读题面,只质疑不改写 ----
    critic_prompt = (
        "你是数学建模方案的独立质疑者,任务是**只质疑、不改写**。\n"
        "你独立阅读题目与建模初稿,找出其中真正的问题。\n"
        f"【题目】{problem_str}\n"
        f"【各小问】{problem_index}\n"
        f"【建模分析】{analysis}\n"
        f"【建模初稿 plan_struct】\n{draft}\n\n"
        "请从以下四个角度逐条质疑,只输出质疑清单,不要修改方案:\n"
        "1. 分类质疑:这题真的属于该 problem_type 吗?给出证据;\n"
        "2. 假设风险:哪些假设不成立、遗漏或过于理想;\n"
        "3. 目标函数/约束:可解性、量纲、缺失约束;\n"
        "4. 方法集建议:若分类错误,候选方法集应换成什么范式组合(机理/数据/仿真)。\n"
        "格式:最后用 ```json 代码围栏输出质疑清单 JSON:\n"
        '{"criticisms": [{"type": "分类|假设|目标|方法", "issue": "...", "evidence": "..."}]}\n'
        "围栏后不要再输出任何文字。"
    )
    criticisms = []
    try:
        critic_resp = _invoke_llm(model_text, [HumanMessage(critic_prompt)], action="辩论·质疑者")
        try:
            plan_json, _ = _extract_last_json_fence(critic_resp.content)
            criticisms = plan_json.get("criticisms") or []
        except Exception:
            criticisms = []
    except Exception as e:
        logger.warning(f'debate_plan 质疑者调用失败,静默降级: {e}')
        return {}
    if not criticisms:
        logger.info('debate_plan: 质疑者未提出有效质疑,采用初稿继续修订')

    # ---- 2. 修订者:逐条回应,产出修订版方案+方法卡+分歧点 ----
    critic_block = "\n".join(
        f"- [{c.get('type', '')}] {c.get('issue', '')} (证据: {c.get('evidence', '')})"
        for c in criticisms[:10] if isinstance(c, dict) and c.get('issue')
    ) or "(质疑者未提出具体质疑)"
    methods_line = "、".join(str(m) for m in initial_methods) or "、".join(METHODS_FALLBACK)
    reviser_prompt = (
        "你是数学建模方案的修订者。建模初稿被独立质疑者攻击,请逐条回应并修订。\n"
        "规则:合理的质疑必须采纳;不合理的可以拒绝,但必须在 disagreements 里记录理由。\n"
        f"【题目】{problem_str}\n"
        f"【建模分析】{analysis}\n"
        f"【初版 plan_struct】\n{draft}\n"
        f"【质疑清单】\n{critic_block}\n"
        f"【初版候选方法】{methods_line}\n\n"
        "输出(```json 围栏包裹,围栏后不要输出任何文字):\n"
        "{\n"
        '  "plan_struct": { 修订后的方案,字段与初稿一致: problem_type/variables/objective/constraints/per_question_method },\n'
        '  "methods": [ 方法卡数组,3个,覆盖不同范式(机理/数据/仿真),每项: {"name":"方法名", "paradigm":"范式", "rationale":"为什么用", "assumption":"该方法独立假设", "tools":"典型算法"} ],\n'
        '  "disagreements": [ {"issue":"争议点", "adopter":"质疑者|修订者", "status":"未解决|已解决", "note":"说明"} ]\n'
        "}"
    )
    rev_plan, rev_methods, rev_dis = {}, [], []
    try:
        reviser_resp = _invoke_llm(model_text, [HumanMessage(reviser_prompt)], action="辩论·修订者")
        try:
            rev_json, _ = _extract_last_json_fence(reviser_resp.content)
            rev_plan = rev_json.get("plan_struct") or {}
            rev_methods = rev_json.get("methods") or []
            rev_dis = rev_json.get("disagreements") or []
        except Exception as e:
            logger.warning(f'debate_plan 修订者输出解析失败,静默降级: {e}')
            return {}
        if not rev_plan or not isinstance(rev_methods, list) or not rev_methods:
            logger.warning('debate_plan 修订输出不完整(缺 plan_struct 或 methods),静默降级')
            return {}
    except Exception as e:
        logger.warning(f'debate_plan 修订者调用失败,静默降级: {e}')
        return {}

    # ---- 3. 收口:修订版方案强制结构化(格式硬保证),失败则用修订版原文 ----
    try:
        struct_out = _invoke_llm(model_plan_struct, [
            HumanMessage(
                "把下面的建模方案整理成严格符合 schema 的结构化输出,保持所有数值与内容不变:\n"
                + json.dumps(rev_plan, ensure_ascii=False)
            )
        ], action="辩论·收口")
        plan_struct = json.dumps(dict(struct_out), ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f'debate_plan 收口失败,采用修订版原文: {e}')
        plan_struct = json.dumps(rev_plan, ensure_ascii=False, indent=2)

    # 方法卡兜底:过滤字段不完整的,只保留合法项
    cleaned = []
    for m in rev_methods[:3]:
        if not isinstance(m, dict) or not m.get("name"):
            continue
        cleaned.append({
            "name": str(m.get("name")),
            "paradigm": str(m.get("paradigm") or "通用"),
            "rationale": str(m.get("rationale") or ""),
            "assumption": str(m.get("assumption") or ""),
            "tools": str(m.get("tools") or ""),
        })
    if not cleaned:
        logger.warning('debate_plan 方法卡全部无效,静默降级用初版方法集')
        return {}

    disagreements = [
        {"issue": str(d.get("issue", "")), "adopter": str(d.get("adopter", "修订者")),
         "status": str(d.get("status", "未解决")), "note": str(d.get("note", ""))}
        for d in rev_dis if isinstance(d, dict) and d.get("issue")
    ][:10]

    logger.info(f'debate_plan 完成: 修订 plan_struct, 方法卡 {len(cleaned)} 张, 分歧 {len(disagreements)} 条')
    return {
        "plan_struct": plan_struct,
        "methods": cleaned,
        "disagreements": disagreements,
    }

#执行代码用的 Python 解释器:默认跟随当前运行环境(与 agent 同一解释器),可用环境变量 MATH_PYTHON_EXE 覆盖
PYTHON_EXE = os.environ.get("MATH_PYTHON_EXE") or sys.executable
# 代码/图片目录随当前题目工作区动态解析（多题隔离）
def code_dir() -> Path:
    return ws_root() / "code"

def photo_dir() -> Path:
    return ws_root() / "photo"

# 出图脚本执行:子进程重试次数 + LLM 自动修复轮数上限(均防死循环)
RUN_RETRY = 2
LLM_MAX_FIX = 2
RUN_TIMEOUT = 30
# 纯计算脚本(无 savefig/matplotlib):跑 1 次 + 失败重试 1 次;
# stdout 只保留尾部 STDOUT_TAIL 字符——数值结果集中在最后几行 print,
# 截断防止撑爆汇总节点上下文与 SSE 载荷
CALC_RETRY = 1
STDOUT_TAIL = 1200

#从"问题N"键中提取编号用于排序
def _num_key(k: str) -> int:
    m = re.search(r"\d+", k)
    return int(m.group()) if m else 0

#此节点在扇出前对问题索引做兜底归一化(能真正写回共享状态):
#索引意外为空时把整题当作"问题1",保证 modeling 之后的所有读者
#(屏障统计/求解任务/论文章节)都拿到一致且非空的小问清单
def send_problem_index(state: over_all_state) -> dict:
    logger.info('正在运行 send_problem_index 节点')
    if state.get("problem_index"):
        return {}
    fallback = (state.get("problem_str") or "").strip()
    if not fallback:
        return {}
    logger.warning('problem_index 为空，已把整题作为「问题1」兜底')
    return {"problem_index": {"问题1": fallback}}

#屏障节点:候选方法×全部小问都交卷后才放行到 run_solutions,避免汇总读到部分结果
def collect_branches(state: over_all_state) -> dict:
    problem_index = state.get("problem_index") or {}
    qs = sorted(problem_index.keys(), key=_num_key)
    methods = state.get("methods") or METHODS_FALLBACK
    # 兼容方法卡(dict)与字符串:统一取方法名,与 solve_with_method 的 done_pairs 口径一致
    mnames = [(m.get("name") if isinstance(m, dict) else str(m)) for m in methods if m]
    needed = {f"{mn}|{k}" for mn in mnames for k in qs}
    got = set(state.get("done_pairs") or [])
    if needed and needed <= got:
        return Command(goto="run_solutions", update={})
    return {}

#条件边路径函数:为每个方法生成一个并行任务,每个任务携带全部问题+数据文件名清单(数据内容由代码自行读取)
def dispatch_sends(state: over_all_state) -> list:
    problem_index = state.get("problem_index") or {}
    if not problem_index:
        # send_problem_index 兜底后仍为空 = 连题目原文都没有,确实无内容可解
        return [Send("compare_summarize", {})]
    ddir = dataset_dir()
    dataset_files = "\n".join(sorted(f.name for f in ddir.iterdir() if f.is_file())) if ddir.is_dir() else "(dataset 目录为空)"
    cards = state.get("methods") or METHODS_FALLBACK
    disagreements = state.get("disagreements") or []
    # 兼容两种形态:debate_plan 产出的方法卡(dict)或 modeling 初版的字符串方法名
    sends = []
    for c in cards:
        if isinstance(c, dict) and c.get("name"):
            name, card = str(c["name"]), c
        else:
            name, card = str(c), None
        sends.append(Send("solve_with_method", {
            "method": name,
            "method_card": card,
            "disagreements": disagreements,
            "problem_index": problem_index,
            "modeling_analysis": state.get("modeling_analysis") or "",
            "plan_struct": state.get("plan_struct") or "",
            "dataset_files": dataset_files,
        }))
    return sends

#工作节点:单次调用内依次解决所有小问。method 来自 Send 输入全程有效,
#工具循环也在节点内就地完成,避免分支身份经图节点往返丢失(旧实现曾导致 method 变为"未知方法")
def solve_with_method(state: over_all_state) -> dict:
    method = state.get("method") or "未知方法"
    logger.info(f'正在运行 solve_with_method 节点: {method}')
    method_card = state.get("method_card") or {}
    disagreements = state.get("disagreements") or []
    problem_index = state.get("problem_index") or {}
    results_parts = []
    code_files = []
    failed = []

    for q_key in sorted(problem_index.keys(), key=_num_key):
        q_text = problem_index[q_key]
        card_block = ""
        if isinstance(method_card, dict) and method_card.get("name"):
            card_block = (
                f"【方法卡】\n"
                f"建模范式: {method_card.get('paradigm', '')}\n"
                f"方法依据: {method_card.get('rationale', '')}\n"
                f"你的独立假设(可偏离共享方案之处): {method_card.get('assumption', '')}\n"
                f"可用工具/算法: {method_card.get('tools', '')}\n"
            )
        dis_block = ""
        if disagreements:
            dis_block = "【未决分歧点】(可按下述立场展开,也可给出你的判断):\n" + "\n".join(
                f"- {d.get('issue', '')}" for d in disagreements[:5] if isinstance(d, dict)
            ) + "\n"
        prompt = (
            f"你正在用《{method}》解答数学建模问题。\n"
            f"当前小问: {q_key}\n题干:\n{q_text}\n"
            f"整体建模分析:\n{state.get('modeling_analysis') or '(无)'}\n"
            + card_block
            + dis_block
            + f"结构化建模方案(共享骨架;除上方方法卡假设与分歧点外保持一致):\n{state.get('plan_struct') or '(无)'}\n"
            f"数据集文件清单(dataset 目录):\n{state.get('dataset_files') or '(无)'}\n\n"
            "数据说明: 上方清单中的文件位于本机 dataset/ 目录,"
            "代码中应通过 r\"..\\dataset\\文件名\" 读取真实数据,禁止编造数据。\n"
            "如需查资料或查看工作区已有文件,可调用 search / read_workspace 工具(work_dir 参数必填,只能是 code/paper/photo/dataset 之一,例如 read_workspace(work_dir=\"code\", rel_path=\"q1.py\"));\n"
            "不要调用写入类工具,思路与代码由系统自动保存。\n"
            "请给出该问在本方法下的完整求解,输出两部分:\n"
            "1) 思路与公式: Markdown,含关键建模公式(LaTeX 语法 $...$),不写代码;撰写中文表述时,请先调用 get_writing_skill 获取去 AI 味写作规范并严格遵循;"
            "思路部分同样遵循 Nature 式'问题→方法→结果'简明证据链:先点明本问采用的方法/模型,再展开推导,每段以定量结论或关键公式收束,"
            "避免空话与'首先/其次/最后'式机械连接词,不使用'值得注意的是/综上所述'等套话。\n"
            "2) 求解代码: 用 ```python 代码块包裹的完整可运行 Python 代码。代码规则:\n"
            f"   - 如需画图,必须用 plt.savefig(r\"{method}_{q_key}.png\") 保存到当前目录,文件名以方法名开头避免冲突;\n"
            "   - 画图前必须设置 matplotlib 中文字体: plt.rcParams['font.sans-serif']=['SimHei']; plt.rcParams['axes.unicode_minus']=False。"
        )
        msgs = [HumanMessage(prompt)]
        # 单问全流程隔离：LLM 调用/解析/落盘任一环节失败只损失该问，
        # 记入 failed_qs 交给 feedback_check 处理，绝不拖崩整个分支
        try:
            # 节点内工具循环：模型要调工具就在本地执行并继续，最多 MAX_TOOL_ROUNDS 轮
            for _ in range(MAX_TOOL_ROUNDS + 1):
                resp = _invoke_llm(model_with_tool, msgs, action=f"求解 {method}/{q_key}")
                if not getattr(resp, "tool_calls", None):
                    break
                tool_msgs = []
                for call in resp.tool_calls:
                    tool = TOOLS_BY_NAME.get(call["name"])
                    try:
                        result = tool.invoke(call["args"]) if tool else f"未知工具: {call['name']}"
                    except Exception as e:
                        result = f"工具执行失败: {e}"
                    tool_msgs.append(ToolMessage(content=str(result), tool_call_id=call["id"]))
                msgs = msgs + [resp] + tool_msgs
            else:
                # 工具轮数超限：改用无工具绑定的模型强制输出最终答案
                resp = _invoke_llm(model, [*msgs, HumanMessage(
                    "工具调用次数已达上限。请不要再调用任何工具，直接给出当前小问的完整求解：思路与公式 + ```python 代码块。"
                )], action=f"求解 {method}/{q_key}·工具超限兜底")
                if getattr(resp, "tool_calls", None):
                    raise RuntimeError("工具超限兜底实例仍要求调用工具")

            # 解析最终答案并落盘
            content = resp.content
            m = re.search(r"```(?:[Pp]ython|[Pp]y3?)?\s*\n?(.*?)```", content, re.S)
            if not m:
                raise ValueError("模型未生成可解析的 Python 代码块")
            code = m.group(1).strip()
            idea = re.sub(r"```(?:[Pp]ython|[Pp]y3?)?\s*\n?.*?```", "", content, flags=re.S).strip()

            write_workspace.invoke({"work_dir": "paper", "rel_path": f"{method}_{q_key}_思路.md", "content": idea})
            code_rel = f"{method}_{q_key}_solution.py"
            write_workspace.invoke({"work_dir": "code", "rel_path": code_rel, "content": code})
            code_files.append(code_rel)
            results_parts.append(f"{q_key}: 完成(思路→paper/{method}_{q_key}_思路.md, 代码→code/{code_rel})")
        except Exception as e:
            logger.error(f"[{method}] {q_key} 求解失败: {e}")
            results_parts.append(f"{q_key}: 求解失败({e})")
            failed.append(q_key)

    update = {"answers": [{method: "\n".join(results_parts)}], "code_files": code_files, "failed_qs": failed}
    update["done_pairs"] = [f"{method}|{k}" for k in sorted(problem_index.keys(), key=_num_key)]
    return update

#把 code/ 下所有 png 移入 photo/(同名冲突自动加 _1/_2 后缀),保证 code/ 不残留图片
def _sweep_pngs_to_photo() -> list:
    moved = []
    for img in list(code_dir().rglob("*.png")):
        dest = photo_dir() / img.name
        i = 1
        while dest.exists():
            dest = photo_dir() / f"{img.stem}_{i}{img.suffix}"
            i += 1
        shutil.move(str(img), str(dest))
        moved.append(dest.name)
    return moved

# ---------- 出图脚本执行与自愈 ----------
# 从脚本文本提取 savefig 指定的输出文件名,用于重入时判断图片是否已生成
def _extract_savefig_paths(text: str) -> List[str]:
    return [Path(m.group(1)).name for m in re.finditer(r"savefig\s*\(\s*[rR]?['\"]([^'\"]+)['\"]", text)]


# 执行单个出图脚本一次,返回结构化结果(不移动图片)
def _run_one_script(script: Path, timeout: int) -> dict:
    before = set(code_dir().rglob("*.png"))
    try:
        proc = subprocess.run(
            [PYTHON_EXE, script.name],
            cwd=str(code_dir()),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
            env={**os.environ, "PYTHONIOENCODING": "utf-8", "MPLBACKEND": "Agg"},
        )
        return {
            "returncode": proc.returncode,
            "stdout": proc.stdout or "",
            "stderr": proc.stderr or "",
            "timed_out": False,
            "new_imgs": list(set(code_dir().rglob("*.png")) - before),
            "exc": None,
        }
    except subprocess.TimeoutExpired:
        return {"returncode": None, "stdout": "", "stderr": "", "timed_out": True, "new_imgs": [], "exc": None}
    except Exception as e:
        return {"returncode": None, "stdout": "", "stderr": "", "timed_out": False, "new_imgs": [], "exc": e}


# 把新生成的 png 从 code/ 移入 photo/(同名冲突自动加后缀),返回移动后的文件名
def _move_imgs_to_photo(imgs) -> List[str]:
    moved = []
    for img in imgs:
        dest = photo_dir() / img.name
        i = 1
        while dest.exists():
            dest = photo_dir() / f"{img.stem}_{i}{img.suffix}"
            i += 1
        shutil.move(str(img), str(dest))
        moved.append(dest.name)
    return moved


# 把单次执行结果转成人类可读的错误摘要,成功时返回空串
def _run_error_msg(r: dict, timeout: int) -> str:
    if r["timed_out"]:
        return f"运行超时({timeout}秒),已终止"
    if r["exc"] is not None:
        return f"运行出错: {r['exc']}"
    if r["returncode"] != 0:
        return f"运行失败(返回码 {r['returncode']})"
    return ""


# 脚本运行输出留档: 图会存进 photo/, 但 stdout 里的数值结果此前只活在内存的 run_report
# (且每脚本仅保留末尾 1200 字符), 程序一关数据就丢。这里把完整 stdout/stderr 落盘到
# code/logs/{脚本名}.log——数值与图一并留档, 且不受截断, 供复核与支撑材料使用。
def _save_run_log(rel: str, r: dict) -> None:
    try:
        log_dir = code_dir() / "logs"
        log_dir.mkdir(parents=True, exist_ok=True)
        body = (f"# {rel} 运行输出留档\n"
                f"# 时间: {time.strftime('%Y-%m-%d %H:%M:%S')}\n"
                f"# returncode: {r['returncode']}  timed_out: {r['timed_out']}\n"
                f"--- stdout ---\n{r['stdout'] or '(无)'}\n"
                f"--- stderr ---\n{r['stderr'] or '(无)'}\n")
        (log_dir / f"{Path(rel).stem}.log").write_text(body, encoding="utf-8")
    except Exception as e:
        logger.warning(f"运行输出留档失败({rel}): {e}")


# 提交 LLM 诊断并重写脚本:成功写回磁盘返回 True,否则 False
def _llm_fix_script(script: Path, error_log: List[str], expected: List[str]) -> bool:
    try:
        code = script.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"读取待修复脚本失败: {e}")
        return False
    log_text = "\n".join(error_log)[-2500:]
    exp_hint = f"；期望生成的图片文件名: {expected}" if expected else ""
    prompt = (
        "你是一名 Python 绘图排障专家。下面是一段用于数学建模出图的 Python 脚本，运行时失败了。\n"
        "请先诊断失败原因（重点排查：中文字体缺失/乱码、模块未安装、数据读取路径错误、"
        "数据集文件不存在、数组维度/除零、死循环或超时等），然后直接返回修正后的【完整可运行】代码。\n\n"
        f"错误日志（含返回码/stdout/stderr/是否超时）：\n{log_text}\n\n"
        f"原始脚本：\n{code}\n\n"
        "要求：\n"
        "1. 必须用 ```python 代码块包裹完整代码，不要只给片段；\n"
        "2. 必须保留原有的 plt.savefig(...) 调用，且文件名保持不变" + exp_hint + "，否则图片无法被收集；\n"
        "3. 不要改变数据读取路径（数据在 dataset/ 下，使用相对路径 ..\\dataset\\文件名）；\n"
        "4. 不要引入当前环境未安装的额外依赖；\n"
        "5. 中文字体问题可降级为 plt.rcParams['font.sans-serif']=['Microsoft YaHei'] 或改用英文标签。\n"
        "只输出代码块，不要多余解释。"
    )
    try:
        resp = _invoke_llm(model_text, [HumanMessage(prompt)], action="求解代码·失败诊断修复")
    except Exception as e:
        logger.error(f"LLM 修复调用失败: {e}")
        return False
    m = re.search(r"```(?:[Pp]ython|[Pp]y3?)?\s*\n?(.*?)```", resp.content or "", re.S)
    if not m:
        return False
    new_code = m.group(1).strip()
    if not new_code:
        return False
    try:
        script.write_text(new_code, encoding="utf-8")
        return True
    except Exception as e:
        logger.error(f"写回修复脚本失败: {e}")
        return False


# 执行节点:只运行"负责生成图片"的脚本(静态检测 savefig/matplotlib);
# 单个脚本失败自动重试 RUN_RETRY 次,仍失败则提交 LLM 诊断并改写代码(最多 LLM_MAX_FIX 轮),
# 全部失败才中断交由人工处理。重入时跳过已生成图片的脚本,避免重复执行。
def run_solutions(state: over_all_state) -> dict:
    logger.info('正在运行 run_solutions 节点')
    code_files = state.get("code_files") or []
    reports = []
    code_dir().mkdir(parents=True, exist_ok=True)
    photo_dir().mkdir(parents=True, exist_ok=True)

    # 先清理上次运行/中断遗留的 png,避免"非本次新增"的图片永远移不进 photo/
    leftover = _sweep_pngs_to_photo()
    if leftover:
        reports.append({"file": "(清理遗留)", "status": f"移入上次遗留图片: {leftover}"})

    to_run = []    # 出图脚本:重试 + LLM 修复 + 图片收集
    to_calc = []   # 纯计算脚本:全量执行,捕获 stdout 数值结果
    for rel in code_files:
        script = code_dir() / rel
        if not script.is_file():
            reports.append({"file": rel, "status": "文件不存在,跳过"})
            continue
        text = script.read_text(encoding="utf-8", errors="ignore")
        if "savefig" in text or "matplotlib" in text:
            to_run.append((rel, script))
        else:
            to_calc.append((rel, script))

    # 纯计算脚本:真正执行并记录数值输出——汇总与论文的"关键数值"由此获得真实来源。
    # 失败只记录不进人工介入循环(那是出图脚本专属),不阻塞主流程
    for rel, script in to_calc:
        r = None
        for attempt in range(1 + CALC_RETRY):
            timeout = RUN_TIMEOUT if attempt == 0 else RUN_TIMEOUT * 2
            r = _run_one_script(script, timeout)
            if not _run_error_msg(r, timeout):
                break
        if r is None:
            continue
        timeout = RUN_TIMEOUT
        _save_run_log(rel, r)
        msg = _run_error_msg(r, timeout)
        if msg:
            reports.append({"file": rel, "status": f"失败: {msg}",
                            "error": (r["stderr"] or "")[-400:]})
        else:
            reports.append({"file": rel, "status": "成功",
                            "output": (r["stdout"] or "")[-STDOUT_TAIL:] or "(无输出)"})

    failed = []  # (rel, script, error_log)

    for rel, script in to_run:
        expected = _extract_savefig_paths(script.read_text(encoding="utf-8", errors="ignore"))
        # 重入跳过:期望图片已存在说明上一轮已成功
        if expected and all((photo_dir() / n).exists() for n in expected):
            reports.append({"file": rel, "status": f"成功(已有图片,跳过): {expected}"})
            continue

        error_log: List[str] = []
        success = False
        moved: List[str] = []

        # ① 子进程重试:1 次 + RUN_RETRY 次(重试时超时翻倍)
        for attempt in range(1 + RUN_RETRY):
            timeout = RUN_TIMEOUT if attempt == 0 else RUN_TIMEOUT * 2
            r = _run_one_script(script, timeout)
            msg = _run_error_msg(r, timeout)
            if not msg:
                moved = _move_imgs_to_photo(r["new_imgs"])
                if moved:
                    success = True
                    reports.append({"file": rel, "status": f"成功,生成图片: {moved}（第{attempt + 1}次执行）",
                                    "output": (r["stdout"] or "")[-STDOUT_TAIL:]})
                    break
                msg = "运行成功但未生成任何图片(savefig 未生效?)"
            error_log.append(
                f"[第{attempt + 1}次执行] {msg}"
                + (f" | stdout: {r['stdout'][:400]}" if r["stdout"] else "")
                + (f" | stderr: {r['stderr'][:400]}" if r["stderr"] else "")
            )

        # ② LLM 审查 + 改写代码(最多 LLM_MAX_FIX 轮,每轮重跑一次)
        fix_i = 0
        while not success and fix_i < LLM_MAX_FIX:
            fix_i += 1
            logger.warning(f"[{rel}] 子进程重试失败,第 {fix_i} 次提交 LLM 诊断修复")
            if not _llm_fix_script(script, error_log, expected):
                error_log.append(f"[LLM 修复第{fix_i}轮] 未能解析出可运行代码块,放弃")
                break
            for attempt in range(1 + 1):
                timeout = RUN_TIMEOUT if attempt == 0 else RUN_TIMEOUT * 2
                r = _run_one_script(script, timeout)
                msg = _run_error_msg(r, timeout)
                if not msg:
                    moved = _move_imgs_to_photo(r["new_imgs"])
                    if moved:
                        success = True
                        reports.append({"file": rel, "status": f"成功,生成图片: {moved}（LLM 修复第{fix_i}轮后）",
                                        "output": (r["stdout"] or "")[-STDOUT_TAIL:]})
                        break
                    msg = "运行成功但未生成图片"
                error_log.append(
                    f"[LLM 修复第{fix_i}轮·第{attempt + 1}次执行] {msg}"
                    + (f" | stderr: {r['stderr'][:400]}" if r["stderr"] else "")
                )
            if success:
                break

        _save_run_log(rel, r)   # 无论成败, 完整输出留档 code/logs/
        if not success:
            failed.append((rel, script, error_log))
            reports.append({"file": rel, "status": f"失败:经 {RUN_RETRY} 次重试 + {LLM_MAX_FIX} 轮 LLM 修复仍未能出图"})

    # 收尾清扫:即使脚本失败/超时,已 savefig 落盘的 png 也一并移入 photo/
    remaining = _sweep_pngs_to_photo()
    if remaining:
        reports.append({"file": "(收尾清扫)", "status": f"移入图片: {remaining}"})

    # ③ 终极兜底:仍有失败 → 中断交人工,可多轮介入直到成功或放弃。
    # 原理:interrupt 的 resume 会让本节点从头重跑——上方的自动恢复流程(含人工刚改的脚本)会再执行一遍;
    # 若重跑后仍失败,下方 while 会再次 interrupt 询问,人工可继续修改再试,直到成功或输入 skip。
    while failed:
        detail = "\n\n".join(f"脚本: {rel}\n" + "\n".join(log) for rel, _, log in failed)
        feedback = interrupt(
            f"以下出图脚本经 {RUN_RETRY} 次重试 + {LLM_MAX_FIX} 轮 LLM 修复仍失败:\n\n"
            f"{detail}\n\n"
            "【回车 = 跳过失败脚本,继续后续流程】\n"
            "【输入 重试/继续 = 重跑(可先修改 code/ 下对应脚本,会带着修改重跑,仍失败会再次询问)】"
        )
        fb = (feedback or "").strip().lower()
        # 语义:空输入或"跳过/通过"类词 = 跳过失败项继续;只有明确要求重试才重跑
        skip = not fb or any(k in fb for k in ("skip", "pass", "通过", "跳过", "略过"))
        if skip:
            for rel, _, _ in failed:
                reports.append({"file": rel, "status": "已由人工选择跳过"})
            log_ai_event("人工审查", "出图失败人工处置: 选择跳过失败脚本")
            break
        # 明确要求重试:本轮 resume 的重跑已在上方执行过;若仍失败则回到循环顶部再次询问
        logger.info(f"人工选择重试(反馈: {fb[:50]}),仍有 {len(failed)} 个失败脚本")
        log_ai_event("人工审查", f"出图失败人工处置: 要求重试(反馈: {fb[:80]})")

    return {"run_report": reports}

#路由:有工具调用→执行工具;工具刚返回→回本节点继续对话;否则→结束
def compare_route(state: over_all_state) -> str:
    msgs = state.get("compare_msgs") or []
    last = msgs[-1] if msgs else None
    if getattr(last, "tool_calls", None):
        return "tools"
    return "continue_next" if isinstance(last, ToolMessage) else "end"

#汇总节点:收集四个方法的答卷(可调工具查看工作区),对比分歧并给出每问最终结论,落盘最终总结
def compare_summarize(state: over_all_state) -> dict:
    logger.info('正在运行 compare_summarize 节点')
    merged = {}
    for item in (state.get("answers") or []):
        for k, v in item.items():
            merged.setdefault(k, []).append(v)
    problem_index = state.get("problem_index") or {}
    run_text = "\n".join(
        f"{r.get('file')}: {r.get('status', '')}"
        + (f" | stdout: {r['output']}" if r.get("output") else "")
        + (f" | stderr: {r['error']}" if r.get("error") else "")
        for r in (state.get("run_report") or [])
    )
    msgs = state.get("compare_msgs") or []
    #工具刚返回:延续对话,让模型看到工具结果后继续
    if msgs and isinstance(msgs[-1], ToolMessage):
        resp = _invoke_llm(model_think_tool, msgs, action="汇总对比·工具续答")
        to_append = []
    else:
        if not merged:
            final = "(候选方法均未产出结果,无法对比总结)"
            write_workspace.invoke({"work_dir": "paper", "rel_path": "最终总结.md", "content": final})
            return {"final_summary": final}
        input_text = "\n\n".join(f"### 方法《{k}》\n" + "\n".join(vs) for k, vs in merged.items())
        method_names = list(merged.keys())
        method_line = "、".join(method_names) if method_names else "(无)"
        fb = (state.get("human_feedback") or "").strip()
        feedback_line = f"人工审核意见(务必据此修改最终总结): {fb}\n\n" if fb else ""
        # 注入各分支范式与未决分歧,让总评知道这是跨范式的独立估计(交叉验证依据)
        paradigm_line = ""
        mc = state.get("methods") or []
        if mc:
            parts = []
            for c in mc:
                if isinstance(c, dict) and c.get("name"):
                    parts.append(f"{c.get('name')}(范式:{c.get('paradigm', '')})")
            if parts:
                paradigm_line = f"各分支范式(交叉验证依据): {'、'.join(parts)}\n\n"
        dis_line = ""
        dis = state.get("disagreements") or []
        if dis:
            dis_line = "建模阶段未决分歧(总评需裁决或标注待人工确认):\n" + "\n".join(
                f"- {d.get('issue', '')}({d.get('adopter', '')}坚持, {d.get('status', '')})"
                for d in dis[:5] if isinstance(d, dict)
            ) + "\n\n"
        prompt = (
            f"你是数学建模竞赛的总评专家。以下是用 {len(method_names)} 种方法({method_line})解答同一份题目的答卷,以及代码实际运行情况。\n"
            f"题目小问: {list(problem_index.keys())}\n\n"
            f"{paradigm_line}{dis_line}"
            f"{input_text}\n\n"
            f"代码运行情况:\n{run_text or '(无运行报告)'}\n\n"
            f"{feedback_line}"
            "如需查看生成的代码或图片,可调用 read_workspace 工具(work_dir 参数必填,只能是 code/paper/photo/dataset 之一)。\n"
            "【数值真实性铁律】代码运行情况中各脚本的 stdout 输出是唯一可信的数值来源;"
            "最终结论里的每一个关键数值都必须能在运行输出或各方法答卷的求解结果中找到出处,"
            "严禁自行编造、估计或'合理化'数值;若某问缺少可信数值,如实写'该问数值待运行确认',不要虚构。\n"
            "请输出一份对比总结 Markdown 文档:\n"
            "1. 逐问对比各方法的结论,标出结论一致与分歧之处;\n"
            "2. 结合代码运行情况(成功/失败/生成的图),说明结果可信度;\n"
            "3. 对每一问给出最可信的最终结论(可综合多种方法),并简要说明理由;\n"
            "4. 整篇文档结构完整,可直接作为《最终总结》;\n"
            "5. 对每一问明确标注:最终采用的方法(从上述方法中选定)+一句话理由+该问关键数值结果清单"
            "(供论文撰写章节直接引用,数值必须与前述结论一致)。\n\n"
            f"【中文写作规范(去 AI 味,务必严格遵循)】\n{read_writing_skill('de-AI')}"
        )
        resp = _invoke_llm(model_think_tool, [HumanMessage(prompt)], action="汇总对比总结")
        to_append = [HumanMessage(prompt)]

    #模型想调工具:交给工具节点,工具结果会自动回到本节点继续
    if getattr(resp, "tool_calls", None):
        return Command(goto="compare_tool_node", update={"compare_msgs": to_append + [resp]})

    final = resp.content
    write_workspace.invoke({"work_dir": "paper", "rel_path": "最终总结.md", "content": final})
    return {"final_summary": final, "compare_msgs": [resp], "human_feedback": ""}

#此函数用来人工审核最终结果
def final_analysis(state:over_all_state) ->Command:
    final_summary=state['final_summary']
    feedback=interrupt(
        f'请人工最终审查\n'
        f'认可:请输入回车\n'
        f'不认可请输入:打回或提出自己的意见\n'
        f'模型最终结果是:{final_summary}'
    )
    feedback=(feedback or '').strip()
    if not feedback or any(k in feedback for k in ('通过','认可')):
        log_ai_event("人工审查", "最终总结人工审查: 认可")
        return Command(goto="fill_document_meta")
    log_ai_event("人工审查", f"最终总结人工审查: 打回修改, 意见: {feedback[:120]}")
    return Command(goto='compare_summarize', update={
        'human_feedback': feedback,
        'compare_msgs': [RemoveMessage(id=m.id) for m in (state.get("compare_msgs") or []) if getattr(m, "id", None)],
    })
    

# ============ 论文撰写节点：把最终结论/思路/图片填入 LaTeX 国赛模板的独立副本 ============
# 输出目录：paper/latex/（随当前题目工作区，独立副本，不污染原始 writter_struct 模板）
def latex_out() -> Path:
    return ws_root() / "paper" / "latex"

WRITER_TEMPLATE = WORKSPACE_ROOT / "writter_struct"
# 模板里只复制一次、之后不再覆盖的静态文件（类文件/编译脚本/参考文献库等）
LATEX_STATIC = ["cumcmthesis.cls", "book.bib", "build.bat", "clean.bat",
                ".gitignore", "LICENSE", "README.md", "常用LaTex代码指令.txt", "document.tex"]

# 各章节：(文件名, 写作要求)，文件名对应 writter_struct/texfile/<name>.tex
ARTICLE_CHAPTERS = [
    ("1abstract",
     "摘要：300~600字，严格按 Nature 式'背景/动机→问题→方法/模型→结果(定量)→意义'证据链组织，"
     "每段以结论或定量结果开头（结果先行），杜绝空话与'首先/其次/最后'式机械连接词。"
     "首段一句话点出研究背景与动机；随后按'针对问题一/问题二/...'逐问成段，每段先写该问最终采用的"
     "方法（从【候选方法清单】中选定的最适一种，只写该方法名与所用模型）与所用模型，"
     "紧接着给出关键定量结果（必须写具体数字，如'最优组合(20,40)、最大利润220元'，亮点突出）；"
     "末段一句话写优化推广（该建模思路在何类问题可迁移），不展开。"
     "结尾用 \\keywords{关键词1\\quad 关键词2...}（3~5个，必须含模型名与方法名亮点）。只接受文字，不要图表。"),
    ("2ProblemRestatement",
     "问题重述：\\section{问题重述}，含 \\subsection{问题背景}（提炼原题，保持原意）与 "
     "\\subsection{问题提出}（用 enumerate 列出各小问，基本原样复制题目所问）。篇幅不超一页。"),
    ("3ProblemAnalysis",
     "问题分析：\\section{问题分析}，对各小问分别 \\subsection{问题X分析} 写数据洞察与建模思路，"
     "对应 problem_index 的每个小问；并在每问分析末用一句话自然给出'本问拟采用的最适方法及理由'"
     "（从【候选方法清单】中选定，依据最终总结中的对比），为后文建模章节定主线。"
     "要求：数据洞察尽量带具体数字（如'果汁单瓶利润4元、利润率100%'）；方法理由用逻辑自然衔接，"
     "避免'首先/其次/最后'式套话；整体遵循 Nature 式'问题→方法→结果'简明证据链，不写空话。"),
    ("4AssumptionAndSign",
     "模型假设与符号说明：\\section{模型的假设} 用 enumerate 列 5 条左右合理假设；"
     "\\section{符号说明} 用 booktabs 三线表（符号|说明）。"),
    ("5MakeModel",
     "模型的建立与求解：\\section{模型的建立与求解}。先写数据预处理，再逐问 "
     "\\subsection{问题X的模型建立与求解}：先一句话声明'本问最终采用的方法及其理由'"
     "（以最终总结中选定的最适方法为主线，其余方法至多一句话对比），再按 模型建立/求解/结果 三子节展开；"
     "结果子节必须给出与最终总结一致的具体数值，并'结果先行'——先抛结论性定量结果（如'最优解为(20,40)，最大利润220元'），再简述如何得到；"
     "可引用'可用图片'中与本问内容匹配的图"
     "\\includegraphics{图片名.png}（graphicspath 已指向 texfile/figures/，直接写裸文件名即可，"
     "不要加目录前缀）并配 \\caption/\\label（每问建议1~2张，图注须自包含、直接说明图中反映了什么结果/趋势）；"
     "公式用 $...$ 或 equation 环境，关键公式必须完整可推导；所有数学命令（\\mathrm 等）只能出现在公式内，禁止在正文文本中使用；"
     "整体遵循 Nature 式'方法→结果→讨论'简明证据链，段落间用逻辑衔接，避免机械连接词堆砌。"),
    ("6ErrorAnalysis",
     "误差分析：\\section{误差分析}，逐问 \\subsection{针对问题X的误差分析} 说明结果检验与误差来源。"
     "要求：检验结果尽量定量（如'与解析解偏差≤X%'、'交叉验证误差为X'），避免'结果较为可靠'等空泛表述；"
     "误差来源点到关键项即可，遵循 Nature 式简明证据链。"),
    ("7ModelEvaluation",
     "模型评价：\\section{模型的评价}，含 优点/缺点/推广 三 subsection，用 itemize 罗列。"
     "要求：优点/缺点各写具体（结合本模型实际，避免'模型简单有效'等套话）；"
     "'推广'须写成可迁移性——指出该建模思路/方法在哪些同类问题可复用、需如何调整，而非空泛一句；"
     "整体遵循 Nature 式'意义/讨论'导向。"),
    ("8Reference",
     "参考文献：用 thebibliography 环境（\\bibitem{ref01}...），按 GB/T 7714 风格列出正文中实际引用文献。"
     "AI 工具使用声明由系统单独成节（置于本节之前），本节不要写任何 AI 声明内容。"),
    ("9Appendix",
     "附录：\\appendix。\\section{详细图表} 索引；\\section{代码程序} 用 "
     "\\lstinputlisting[style=Python, caption={...}]{code/文件名.py} 列出'求解代码文件'中每个 py；"
     "\\section{支撑材料} 枚举。"),
]


def _ensure_latex_skeleton():
    """首次把 writter_struct 模板静态文件复制到 paper/latex 独立副本；已存在则不覆盖。"""
    latex_out().mkdir(parents=True, exist_ok=True)
    for name in LATEX_STATIC:
        src = WRITER_TEMPLATE / name
        dst = latex_out() / name
        if src.exists() and not dst.exists():
            shutil.copy(src, dst)
    (latex_out() / "texfile").mkdir(parents=True, exist_ok=True)
    (latex_out() / "texfile" / "figures").mkdir(parents=True, exist_ok=True)
    (latex_out() / "code").mkdir(parents=True, exist_ok=True)
    # 复制模板自带的 code 示例，避免附录 \lstinputlisting 缺文件
    tcode = WRITER_TEMPLATE / "code"
    if tcode.is_dir():
        for f in tcode.iterdir():
            if f.is_file() and not (latex_out() / "code" / f.name).exists():
                shutil.copy(f, latex_out() / "code" / f.name)


def _strip_tex_fences(text: str) -> str:
    """去掉 LLM 常包裹的 ```latex ... ``` 代码围栏，避免原样写进 .tex 导致编译失败。"""
    s = (text or "").strip()
    if s.startswith("```"):
        s = re.sub(r"^```[^\n]*\n", "", s)
        s = re.sub(r"\n```\s*$", "", s)
    return s.strip()


def _clean_latex_output():
    """清空上一轮生成的章节/图片/代码，避免换题或重跑时残留旧产物。"""
    for f in (latex_out() / "texfile").glob("*.tex"):
        f.unlink()
    for d in ((latex_out() / "texfile" / "figures"), (latex_out() / "code")):
        for f in d.glob("*"):
            if f.is_file():
                f.unlink()


# ---------- LaTeX 编译验证与自愈：生成后真实编译，报错喂回 LLM 修复，循环到零错误 ----------
LATEX_MAX_FIX = 3    # 编译报错后 LLM 自动修复的最大轮数
LATEX_TIMEOUT = 180  # 单次 xelatex 编译超时(秒)


def _find_xelatex():
    """定位 xelatex：环境变量 MATH_XELATEX > PATH > C:/texlive/*/bin/windows/。找不到返回 None(跳过编译验证)。"""
    cand = os.environ.get("MATH_XELATEX") or shutil.which("xelatex")
    if cand:
        return cand
    for p in Path("C:/texlive").glob("*/bin/windows/xelatex.exe"):
        return str(p)
    return None


def _run_xelatex(xe: str) -> str:
    """在 latex_out() 跑一遍 xelatex，返回 document.log 文本；先清理残留的 synctex 锁文件。"""
    for busy in latex_out().glob("*.synctex(busy)"):
        try:
            busy.unlink()
        except OSError:
            pass
    try:
        subprocess.run(
            [xe, "-interaction=nonstopmode", "document.tex"],
            cwd=str(latex_out()), capture_output=True, text=True,
            encoding="utf-8", errors="replace", timeout=LATEX_TIMEOUT,
        )
    except Exception as e:
        return f"! xelatex 运行失败: {e}"
    log_path = latex_out() / "document.log"
    return log_path.read_text(encoding="utf-8", errors="replace") if log_path.exists() else ""


def _extract_latex_errors(log: str) -> List[Dict[str, str]]:
    """从 xelatex 日志提取 '!' 开头的错误块，并按日志中最近打开的 .tex 文件归属错误位置。"""
    lines = log.splitlines()
    opens = [(i, m.group(1)) for i, ln in enumerate(lines)
             for m in re.finditer(r"[ (]\./(\S+\.tex)", ln)]
    errors = []
    for i, ln in enumerate(lines):
        if not ln.startswith("!"):
            continue
        fname = ""
        for oi, f in opens:
            if oi <= i:
                fname = f
            else:
                break
        ctx = " | ".join(x.strip() for x in lines[i:i + 8] if x.strip())
        errors.append({"file": fname, "msg": ctx[:600]})
    return errors


def _fix_latex_with_llm(errors: List[Dict[str, str]]) -> List[str]:
    """把编译错误与相关章节文件内容交给 LLM 修复并写回，返回修复的文件名列表。"""
    files = sorted({e["file"] for e in errors if e["file"].startswith("texfile/")})
    if not files:
        logger.warning(f"编译错误未能定位到章节文件: {[e['file'] for e in errors]}")
        return []
    parts = ["以下 CUMCM 论文的 LaTeX 章节编译报错，请修复。"]
    for e in errors:
        parts.append(f"[错误·{e['file']}] {e['msg']}")
    for f in files:
        path = latex_out() / f
        if path.exists():
            parts.append(f"===== 文件 {f} 完整内容 =====\n" + path.read_text(encoding="utf-8", errors="replace"))
    parts.append(
        "修复要求：只修导致编译错误的语法问题（数学环境缺 $、数学命令用在文本模式、括号/环境不配对、"
        "\\includegraphics 文件名错误、非法字符等），保持文字内容不变、不改写论述。\n"
        "输出格式（可包含多个文件，除此之外不要任何解释文字）：\n"
        "===FILE: texfile/文件名.tex===\n修复后的完整文件内容\n===END==="
    )
    try:
        resp = _invoke_llm(model_text, [HumanMessage("\n\n".join(parts))], action="论文·编译错误修复")
    except Exception as e:
        logger.error(f"LLM 修复编译错误调用失败: {e}")
        return []
    fixed = []
    for m in re.finditer(r"===FILE:\s*(\S+)\s*===\n(.*?)\n?===END===", resp.content or "", re.S):
        fname, content = m.group(1).strip(), m.group(2).strip()
        if fname.startswith("texfile/") and content:
            content = _strip_tex_fences(content)
            (latex_out() / fname).write_text(content + "\n", encoding="utf-8")
            fixed.append(fname)
    return fixed


# ---------- 2026 AI 使用规定的合规产物: 声明章节 + 《AI 工具使用详情》 ----------
# 全部程序化拼装,不经 LLM——规定给了逐字模板,合规文本不允许模型自由发挥。

# AI 使用环节分类: (环节名, 动作关键词, 详情说明, 声明中的用途短语)
_AI_CATEGORIES = [
    ("题目信息提取", ("问题提取",), "用结构化提示从题面提取小问索引与完整题干", "题目信息提取"),
    ("建模分析与质检", ("建模分析", "建模质检"), "基于题面与数据集生成建模分析,并对方案做结构化质检", "建模方案分析"),
    ("建模方案辩论", ("辩论",), "质疑者-修订者双角色对建模方案交叉质疑、修订并收敛", "建模方案分析"),
    ("求解代码生成", ("求解",), "按候选方法逐问生成解题思路与可运行 Python 代码", "求解代码生成"),
    ("代码调试修复", ("失败诊断", "诊断修复"), "对运行失败的求解/绘图脚本诊断原因并改写修复", "代码调试"),
    ("结果对比总结", ("汇总对比",), "汇总各方法答卷与代码运行输出,逐问对比并生成最终总结", "结果对比总结"),
    ("论文撰写", ("论文撰写",), "按国赛 LaTeX 模板逐章撰写论文", "论文撰写与语言润色"),
    ("论文编译修复", ("编译错误",), "定位并修复论文 LaTeX 编译错误", "论文撰写与语言润色"),
    ("标题拟定", ("封面标题",), "根据题面拟定论文标题", "标题拟定"),
]

def _tex_escape(s: str) -> str:
    """LaTeX 特殊字符转义(人工意见/提示摘录等任意文本嵌入 .tex 前必须过这里)。"""
    repl = {"\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
            "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
            "~": r"\textasciitilde{}", "^": r"\textasciicircum{}"}
    return "".join(repl.get(c, c) for c in (s or ""))

def _ai_purpose_phrase(events) -> str:
    """从事件日志按环节规则化提炼声明中的'主要用于'短语(只写出现过的环节)。"""
    used = []
    for _name, kws, _desc, phrase in _AI_CATEGORIES:
        if phrase in used:
            continue
        if any(e.get("kind") == "调用" and any(k in e.get("summary", "") for k in kws)
               for e in events):
            used.append(phrase)
    return "、".join(used) if used else "辅助计算与文本整理"

def _build_ai_declaration(events, override_phrase: str = "") -> str:
    """按 2026 规定固定句式拼装《AI 工具使用声明》。events 为空时生成'未使用'版本。"""
    events = events or []
    if not events:
        body = "本参赛队在竞赛过程中未使用任何 AI 工具。"
    else:
        purpose = (override_phrase or "").strip() or _ai_purpose_phrase(events)
        body = (f"本参赛队在竞赛过程中使用了 AI 工具，主要用于{purpose}，"
                "详细使用情况见支撑材料。")
    return "\\newpage\n\n\\section*{AI 工具使用声明}\n\n" + _tex_escape(body) + "\n"

def _ensure_ai_declaration_input() -> None:
    """确保 document.tex 在参考文献前引入 AI 声明章节。
    修补而非只改模板: 老工作区已复制的旧 document.tex 也能被治愈。"""
    doc = latex_out() / "document.tex"
    if not doc.exists():
        return
    text = doc.read_text(encoding="utf-8")
    if "8aAIDeclaration" in text:
        return
    marker = "\\input{texfile/8Reference}"
    if marker in text:
        text = text.replace(
            marker,
            "\\input{texfile/8aAIDeclaration}%AI 工具使用声明(2026 规定,置于参考文献之前)\n" + marker,
            1,
        )
        doc.write_text(text, encoding="utf-8")
        logger.info("document.tex 已插入 AI 声明章节")

def _write_ai_details_pdf() -> str:
    """按 2026 规定生成《AI 工具使用详情.pdf》: 工具版本/使用环节/提示方式/人工核验四节,
    内容全部来自运行事件与模型注册表,不经 LLM。返回给前端的状态描述。"""
    events = ai_events_snapshot()
    out_dir = latex_out() / "ai_details"
    out_dir.mkdir(parents=True, exist_ok=True)

    used_models = []
    for e in events:
        if e.get("kind") == "调用" and e.get("model") and e["model"] not in used_models:
            used_models.append(e["model"])
    if used_models:
        tool_rows = "\n".join(
            f"{_tex_escape(MODEL_REGISTRY.get(k, {}).get('label', k))} & "
            f"{_tex_escape(MODEL_REGISTRY.get(k, {}).get('model', ''))} & OpenAI 兼容 API \\\\"
            for k in used_models)
    else:
        tool_rows = "无 & -- & -- \\\\"

    cat_blocks = []
    for name, kws, desc, _phrase in _AI_CATEGORIES:
        evs = [e for e in events if e.get("kind") == "调用"
               and any(k in e.get("summary", "") for k in kws)]
        if not evs:
            continue
        models = "、".join(sorted({e.get("model", "") for e in evs if e.get("model")})) or "--"
        excerpt = evs[0].get("detail", "")
        excerpt_block = (f"\n\n典型提示摘录（节选）：\\begin{{quote}}\\small {_tex_escape(excerpt)}\\end{{quote}}"
                         if excerpt else "")
        cat_blocks.append(
            f"\\subsection*{{ {_tex_escape(name)} }}\n"
            f"{_tex_escape(desc)}。共调用 {_tex_escape(str(len(evs)))} 次，使用模型：{_tex_escape(models)}。"
            f"{excerpt_block}"
        )
    cat_text = "\n\n".join(cat_blocks) or "本次运行未记录到 AI 调用。"

    human = [e for e in events if e.get("kind") == "人工审查"]
    human_rows = "\n".join(
        f"{_tex_escape(e.get('time', ''))} & {_tex_escape(e.get('summary', ''))} \\\\"
        for e in human) or "无 & -- \\\\"

    now = time.strftime("%Y-%m-%d %H:%M")
    latex_src = f"""\\documentclass[11pt]{{article}}
\\usepackage[UTF8]{{ctex}}
\\usepackage[margin=2.5cm]{{geometry}}
\\usepackage{{booktabs}}
\\begin{{document}}

\\begin{{center}}{{\\LARGE\\bfseries AI 工具使用详情}}\\\\[4pt]
{{\\small 生成时间：{_tex_escape(now)}}}\\end{{center}}

\\section{{所用 AI 工具名称与版本}}
\\begin{{tabular}}{{lll}}
\\toprule
工具 & 模型/版本 & 接入方式 \\\\
\\midrule
{tool_rows}
\\bottomrule
\\end{{tabular}}

\\section{{使用目的与环节}}
全部环节均通过 OpenAI 兼容 API 以中文提示词驱动：各环节先按固定模板构建提示词，调用模型后将输出由程序解析落盘或交人工审核。

{cat_text}

\\section{{人工审查与核验情况}}
系统在流程中设置五类人工确认关卡：建模思路确认、建模分析审核、出图失败处置、最终总结审查、封面信息与声明确认。本次运行记录如下：

\\begin{{tabular}}{{p{{0.22\\textwidth}}p{{0.72\\textwidth}}}}
\\toprule
时间 & 关卡与结果 \\\\
\\midrule
{human_rows}
\\bottomrule
\\end{{tabular}}

\\medskip
本详情所述内容均为系统自动留存的运行记录，即本次运行中对 AI 输出的采纳、人工修改与核验的主要情况；参赛队据此对 AI 参与完成的内容逐项人工审查与核实，核心建模与分析由参赛队主导完成。

\\end{{document}}
"""
    tex = out_dir / "ai_usage_details.tex"
    tex.write_text(latex_src, encoding="utf-8")
    xe = _find_xelatex()
    if not xe:
        return "未检测到 xelatex，《AI 工具使用详情》仅生成源文件(paper/latex/ai_details/)"
    try:
        subprocess.run([xe, "-interaction=nonstopmode", "ai_usage_details.tex"],
                       cwd=str(out_dir), capture_output=True, text=True,
                       encoding="utf-8", errors="replace", timeout=180)
    except Exception as e:
        return f"AI 详情 PDF 编译失败: {e}(源文件在 paper/latex/ai_details/)"
    pdf = out_dir / "ai_usage_details.pdf"
    if pdf.exists():
        shutil.copy(str(pdf), str(ws_root() / "AI 工具使用详情.pdf"))
        return "《AI 工具使用详情.pdf》已生成于工作区根目录"
    return "AI 详情 PDF 编译失败(源文件在 paper/latex/ai_details/)"


def write_article(state: over_all_state) -> dict:
    logger.info("正在运行 write_article 节点：生成 LaTeX 论文到 paper/latex")
    _ensure_latex_skeleton()
    _clean_latex_output()

    final_summary = state.get("final_summary") or "(无最终总结)"
    approach = (state.get("modeling_approach") or state.get("modeling_analysis") or "(无思路)")
    plan_struct = state.get("plan_struct") or ""
    models = state.get("methods") or METHODS_FALLBACK
    # methods 可能是方法卡(dict)或方法名(str),统一取名字
    candidate_methods = "、".join(
        (m.get("name") if isinstance(m, dict) else str(m)) for m in models if m
    ) or "、".join(METHODS_FALLBACK)
    problem_str = state.get("problem_str") or ""
    problem_index = state.get("problem_index") or {}
    code_files = state.get("code_files") or []

    # 1. 同步图片：photo/*.png -> paper/latex/texfile/figures/
    figs = []
    for img in photo_dir().glob("*.png"):
        shutil.copy(img, latex_out() / "texfile" / "figures" / img.name)
        figs.append(img.name)
    # 2. 同步求解代码：code/*.py -> paper/latex/code/（供附录 \lstinputlisting 引用）
    for cf in code_files:
        src = code_dir() / cf
        if src.exists():
            shutil.copy(src, latex_out() / "code" / cf)

    skill = read_writing_skill("de-AI")
    # 汇集求解阶段全部思路与推导(12 份方法答卷的真实运行产物), 原文喂给各章生成:
    # 不设篇幅预算、不加深度清单——素材给足, 章节厚度由模型按素材自行发挥。
    idea_texts = []
    try:
        for f in sorted((ws_root() / "paper").glob("*_思路.md")):
            t = f.read_text(encoding="utf-8", errors="replace").strip()
            if t:
                idea_texts.append(f"### {f.name}\n{t}")
    except Exception as e:
        logger.warning(f"思路素材汇集失败(不影响主流程): {e}")
    idea_block = "\n\n".join(idea_texts)
    done = []
    for fname, req in ARTICLE_CHAPTERS:
        prompt = (
            "你是数学建模国赛论文撰写专家，请把以下内容写成符合 CUMCM LaTeX 模板的【单章节】LaTeX 源码。\n"
            "【写作总则·务必遵循】\n"
            "本论文采用 Nature 学术风：以'背景/动机→问题→方法/模型→结果(定量)→意义/讨论'为统一证据链，"
            "文风简明、直接、证据驱动。它与下方【中文写作规范(去 AI 味)】互补——去 AI 味负责拔除'路标词/讲义腔'等坏毛病，"
            "本总则负责立好文风，二者同时生效、互不替代。\n"
            "1. 结果先行：每个小节/段落先抛出结论或定量结果，再展开方法与推导，不绕弯子；\n"
            "2. 证据驱动：任何判断都必须带具体数字、公式或图引用作支撑，禁止空话、套话、无依据的断言；\n"
            "3. 方法主线：各小问以【最终结论】中选定的最适方法为主线展开（只写一种主线方法，其余方法至多一句话对比），"
            "全文方法口径与最终总结保持一致；\n"
            "4. 图表引用：仅引用文件名与本章内容明显匹配的图，用 \\includegraphics{裸文件名.png} "
            "（graphicspath 已指向 texfile/figures/，禁止加目录前缀），并为每张图配自包含 \\caption——"
            "图注直接说明'该图反映了什么结果/趋势'，可独立读懂，宁缺毋滥；\n"
            "5. 讨论落到意义：模型评价与误差分析不只罗列条目，要指出'为什么重要、局限何在、在何类问题可推广'；\n"
            "6. 数值一致性：文中所有关键数值必须与【最终结论】一致，不得自行编造或改写；\n"
            "7. 公式规范：所有数学命令（如 \\mathrm、\\sum、\\frac）必须写在 $...$ 或 equation 环境内，"
            "正文文本中严禁出现数学命令；公式前后括号配对完整，禁止缺 $。\n"
            "8. 克制连接词：减少'首先/其次/最后/综上所述/值得注意的是'等机械连接词，用逻辑自然衔接"
            "（本规则覆盖模板注释里的'总分总'要求，但国赛必写项——方法、模型、结果——必须写出）。\n\n"
            f"【本章要求】{req}\n\n"
            f"【题目全文】{problem_str}\n\n"
            f"【各小问题干】{problem_index}\n\n"
            f"【建模思路】{approach}\n\n"
            f"【候选方法清单(论文中'采用的模型/方法'只能从其中选用,与最终总结口径一致)】{candidate_methods}\n\n"
            f"【结构化建模方案(JSON,目标函数/约束/变量可直接用于符号说明与模型建立章节,数值须与其他章节一致)】\n{plan_struct or '(无)'}\n\n"
            f"【最终结论（逐问）】{final_summary}\n\n"
            f"【求解阶段完整思路与推导素材（来自全部方法答卷的真实运行产物，写本章时可自由取用、展开与引用，不得虚构其中不存在的推导）】\n{idea_block or '(无)'}\n\n"
            f"【可用图片（位于 texfile/figures/）】{figs or '(无)'}\n"
            f"【求解代码文件（位于 code/，供附录引用）】{code_files or '(无)'}\n\n"
            f"【中文写作规范（去 AI 味，务必遵循）】\n{skill}\n\n"
            "【Nature 风自检清单（输出前逐项核对，不达标则重写）】\n"
            "① 每段是否都有支撑证据（数字/公式/图引用），无空话？\n"
            "② 是否结果先行（先结论/定量结果，再展开）？\n"
            "③ 是否过度使用'首先/其次/最后/综上所述'等机械连接词？\n"
            "④ 所有关键数值是否与【最终结论】一致？\n"
            "⑤ 引用的图其 \\caption 是否自包含、说明了反映的结果？\n"
            "⑥ 讨论/评价是否落到意义、局限与可推广性，而非空泛罗列？\n\n"
            "只输出该章节的完整 LaTeX 正文（不要 \\documentclass、不要 \\begin{document}、不要解释性文字、不要代码围栏），"
            "直接可被 \\input 引用。"
        )
        target = latex_out() / "texfile" / f"{fname}.tex"
        try:
            raw_resp = _invoke_llm(model_text, [HumanMessage(prompt)], action=f"论文撰写·{fname}")
            raw = raw_resp.content
            latex = _strip_tex_fences(raw)
            if not latex:
                raise ValueError("模型返回为空")
            target.write_text(latex, encoding="utf-8")
            done.append(fname)
        except Exception as e:
            logger.error(f"生成章节 {fname} 失败: {e}")
            # 失败也写占位，保证 \\input 不报缺文件
            target.write_text(f"% 本章（{fname}）生成失败：{e}\n", encoding="utf-8")
            done.append(f"{fname}(失败:{e})")

    # 2.5 AI 工具使用声明(2026 规定: 参考文献之前; 固定句式程序化拼装,不经 LLM)
    (latex_out() / "texfile" / "8aAIDeclaration.tex").write_text(
        _build_ai_declaration(ai_events_snapshot()), encoding="utf-8")
    _ensure_ai_declaration_input()
    done.append("8aAIDeclaration")

    # 3. 编译验证与自愈：真实编译一遍，有错误就把日志喂回 LLM 修复，直到零错误或用完修复轮数
    xe = _find_xelatex()
    compile_status = "未检测到 xelatex(可用环境变量 MATH_XELATEX 指定路径)，已跳过编译验证"
    if xe:
        logger.info(f"检测到 xelatex({xe})，开始论文编译验证与自愈")
        log = _run_xelatex(xe)
        errs = _extract_latex_errors(log)
        ok = not errs and "Output written" in log
        for rnd in range(1, LATEX_MAX_FIX + 1):
            if ok:
                break
            logger.warning(f"[编译修复] 第 {rnd}/{LATEX_MAX_FIX} 轮：{len(errs)} 个错误，提交 LLM 修复")
            if not _fix_latex_with_llm(errs):
                break
            log = _run_xelatex(xe)
            errs = _extract_latex_errors(log)
            ok = not errs and "Output written" in log
        if ok:
            _run_xelatex(xe)  # 再跑一遍稳定交叉引用/目录，出最终 PDF
            compile_status = "编译通过(0 错误)，已生成 paper/latex/document.pdf"
        else:
            compile_status = (f"仍有 {len(errs)} 个编译错误(已尝试 {LATEX_MAX_FIX} 轮 LLM 修复)，"
                              f"请人工查看 paper/latex/document.log")
        logger.info(f"论文编译验证结果：{compile_status}")

    details_msg = _write_ai_details_pdf()
    return {"article_chapters": done,
            "compile_status": compile_status,
            "messages": [AIMessage(content="已生成 LaTeX 论文章节：" + "、".join(done)
                                   + f"；{compile_status}；{details_msg}")]}


# 回填 document.tex 封面元信息：标题由 LLM 自动生成，题号/报名号/学校/年份由人工 interrupt 提供
def fill_document_meta(state: over_all_state) -> dict:
    logger.info("正在运行 fill_document_meta 节点：回填 document.tex 元信息")
    doc_path = latex_out() / "document.tex"
    if not doc_path.exists():
        return {"messages": [AIMessage(content="document.tex 不存在，跳过元信息回填")]}

    # 1) 索取需人工确定的元信息（留空=保持模板默认/注释）
    meta = interrupt(
        "请填写论文封面元信息（每行一项，回车留空则保持模板默认）：\n"
        "第1行 题号(如 A/B/C/D)\n"
        "第2行 报名号\n"
        "第3行 学校名称\n"
        "第4行 年份(如 2025)\n"
        "第5行 AI声明用途短语(回车=按运行记录自动生成)\n"
        "直接回车=全部用模板默认"
    )
    lines = (meta or "").splitlines()
    tihao = lines[0].strip() if len(lines) > 0 else ""
    baoming = lines[1].strip() if len(lines) > 1 else ""
    school = lines[2].strip() if len(lines) > 2 else ""
    year = lines[3].strip() if len(lines) > 3 else ""
    ai_phrase = lines[4].strip() if len(lines) > 4 else ""
    # AI 声明按人工确认的用途短语重写(顺带在老工作区补齐声明章节),随元信息回填后的重编译一并生效
    (latex_out() / "texfile" / "8aAIDeclaration.tex").write_text(
        _build_ai_declaration(ai_events_snapshot(), ai_phrase), encoding="utf-8")
    _ensure_ai_declaration_input()
    log_ai_event("人工审查", f"封面信息与 AI 声明确认(用途短语: {ai_phrase or '自动生成'})")

    text = doc_path.read_text(encoding="utf-8")

    # 2) 自动生成标题（基于题目）
    try:
        title_prompt = (
            "请用一句话(中文,不超过30字)概括以下数学建模题目的论文标题，"
            "要求准确、学术、体现核心方法或对象，不要带书名号。只输出标题本身。\n"
            f"题目：{state.get('problem_str') or ''}"
        )
        new_title = _invoke_llm(model, [HumanMessage(title_prompt)], action="封面标题生成").content.strip().strip("《》").strip()
    except Exception as e:
        logger.warning(f"生成标题失败: {e}")
        new_title = ""
    if new_title:
        # 注意:re.sub 的替换字符串会解析 \t \b \s 等转义,直接传 f"\\title{...}" 会把反斜杠吞掉
        # (\t→制表符)甚至抛 re.error。必须用 lambda 返回字面值,避免转义处理。
        text = re.sub(r"\\title\{[^}]*\}", lambda _: f"\\title{{{new_title}}}", text)

    # 3) 回填人工字段（提供则取消注释并填值，保留行尾 %注释；%? 兼容已启用过的行，保证可重复更新）
    if tihao:
        text = re.sub(r"%?\s*\\tihao\{[^}]*\}", lambda _: f"\\tihao{{{tihao}}}", text, count=1)
    if baoming:
        text = re.sub(r"%?\s*\\baominghao\{[^}]*\}", lambda _: f"\\baominghao{{{baoming}}}", text, count=1)
    if school:
        text = re.sub(r"%?\s*\\schoolname\{[^}]*\}", lambda _: f"\\schoolname{{{school}}}", text, count=1)
    if year:
        text = re.sub(r"%?\s*\\yearinput\{[^}]*\}", lambda _: f"\\yearinput{{{year}}}", text, count=1)

    doc_path.write_text(text, encoding="utf-8")

    # 4) 元信息回填后重新编译，保证最终 PDF 包含封面字段(找不到 xelatex 则跳过)
    xe = _find_xelatex()
    final_build = ""
    if xe:
        log = _run_xelatex(xe)
        errs = _extract_latex_errors(log)
        if not errs:
            errs = _extract_latex_errors(_run_xelatex(xe))  # 第二遍稳定引用
        final_build = ("；最终编译通过，PDF 已更新" if not errs
                       else f"；最终编译仍有 {len(errs)} 个错误，请查看 document.log")

    return {"messages": [AIMessage(
        content=f"已回填元信息：题号={tihao or '默认'} 报名号={baoming or '默认'} "
                f"学校={school or '默认'} 年份={year or '默认'} 标题={new_title or '默认'}{final_build}")]}


#测试时建立检查点暂时用内存存储
checkpointer=InMemorySaver()

#Studio/API 模式发起 run 时输入可能为空,单独声明空输入状态避免 EmptyInputError
class input_state(TypedDict):
    pass

builder=StateGraph(state_schema=over_all_state, input=input_state)

builder.add_node('load_problem', load_problem) 
builder.add_node('question_structed',question_structed)
builder.add_node('modeling',modeling)
builder.add_node('debate_plan',debate_plan)
builder.add_node('read_dataset',read_dataset)
builder.add_node('tool_node',tool_node)
builder.add_node('review_modeling_analysis',review_modeling_analysis)
builder.add_node('send_problem_index',send_problem_index)
builder.add_node('collect_branches',collect_branches)
builder.add_node('solve_with_method',solve_with_method)
builder.add_node('feedback_check',feedback_check)
builder.add_node('run_solutions',run_solutions)
builder.add_node('compare_summarize',compare_summarize)
builder.add_node('compare_tool_node',ToolNode(tools=tools, messages_key="compare_msgs"))
builder.add_node('final_analysis', final_analysis)
builder.add_node('write_article', write_article)
builder.add_node('fill_document_meta', fill_document_meta)
builder.add_edge("write_article", "final_analysis")
builder.add_edge("fill_document_meta", END)

builder.add_edge(START, 'load_problem')  
builder.add_edge('load_problem', 'question_structed') 
builder.add_edge('question_structed','read_dataset')
builder.add_edge('read_dataset',"modeling")
builder.add_conditional_edges(
    "modeling",
    modeling_route,
    {"tools": "tool_node", "review": "debate_plan"}
)
builder.add_edge("tool_node", "modeling")
builder.add_edge("debate_plan", "review_modeling_analysis")
builder.add_edge("review_modeling_analysis", "send_problem_index")
builder.add_conditional_edges(
    "send_problem_index",
    dispatch_sends,
    ["solve_with_method", "compare_summarize"]
)
builder.add_edge("solve_with_method", "feedback_check")
builder.add_conditional_edges(
    "compare_summarize",
    compare_route,
    {"tools": "compare_tool_node", "continue_next": "compare_summarize", "end": "write_article"}
)
builder.add_edge("compare_tool_node", "compare_summarize")
builder.add_edge("collect_branches", "run_solutions")
builder.add_edge("run_solutions", "compare_summarize")

# 在 LangGraph API/Server 中运行时，持久化由平台自动管理，不传自定义 checkpointer
if os.environ.get("LANGSMITH_LANGGRAPH_API_VARIANT"):
    graph=builder.compile()
else:
    graph=builder.compile(checkpointer=checkpointer)


if __name__ == "__main__":
    # CLI 直跑：线程 id = 当前题目工作区（默认 default；可用环境变量 MATH_WS 指定）
    ws = os.environ.get("MATH_WS") or get_workspace()
    set_workspace(ws)
    config = {"configurable": {"thread_id": ws}}
    print(f"当前题目: {ws}（工作区: {ws_root()}）")

    #首次运行：执行到第一个 interrupt 处暂停
    graph.invoke({}, config=config)

    #循环处理各中断节点（modeling 思路输入、审核打回等），直到 END
    while True:
        snapshot = graph.get_state(config)
        if not snapshot.next:
            break
        interrupts = getattr(snapshot, "interrupts", None) or []
        prompt = interrupts[0].value if interrupts else "请输入："
        user_input = input(f"{prompt}\n> ")
        graph.invoke(Command(resume=user_input), config=config)

    resp = graph.get_state(config).values
    print(resp)





