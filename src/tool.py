from langchain_core.tools import tool, ToolException
from langgraph.prebuilt import ToolNode
import os
import shlex
import subprocess
import json
from loguru import logger
from pathlib import Path
from dotenv import  load_dotenv
from langchain_community.tools.tavily_search import TavilySearchResults
import pandas as pd

load_dotenv(override=True)

#创建一个搜索的实例
search = TavilySearchResults(max_results=2)

#工作区根目录(由 tool.py 所在位置动态推导,项目迁移后无需改代码)与允许访问的子目录
WORKSPACE_ROOT = Path(__file__).resolve().parent.parent
ALLOWED_DIRS = {"code", "paper", "photo", "dataset"}

# ── 多题目隔离：工作区按题目分目录（workspaces/{题目id}/），set_workspace 切换当前题目 ──
_CURRENT_WS = "default"

def set_workspace(ws_id: str) -> None:
    """切换当前题目工作区（线程串行执行，无并发问题）"""
    global _CURRENT_WS
    _CURRENT_WS = (ws_id or "default").strip() or "default"

def get_workspace() -> str:
    return _CURRENT_WS

def ws_root() -> Path:
    """当前题目的工作区根目录（question/dataset/paper/code/photo 都在其下）"""
    return WORKSPACE_ROOT / "workspaces" / _CURRENT_WS

#解析并校验工作区内的目标路径,非法路径直接抛 ValueError 交由 LangGraph 重试
def _resolve_path(work_dir: str, rel_path: str) -> Path:
    if not work_dir or work_dir not in ALLOWED_DIRS:
        raise ToolException(f"work_dir 参数必填,且只能是 {'/'.join(sorted(ALLOWED_DIRS))} 之一,收到: {work_dir!r}")
    if not rel_path or not rel_path.strip():
        raise ToolException(f"rel_path 不能为空")
    p = Path(rel_path)
    if p.is_absolute() or ".." in p.parts:
        raise ToolException(f"rel_path 必须是相对路径且不能包含 '..',收到: {rel_path}")
    root = ws_root()
    target = (root / work_dir / p).resolve()
    if not target.is_relative_to(root.resolve()):
        raise ToolException(f"路径超出工作区范围: {rel_path}")
    return target

#按后缀智能读取:表格转文本、json 美化、docx 提取,其余按 UTF-8 文本;任何失败统一转 ValueError 交由 LangGraph 重试
def _read_smart(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in (".csv", ".tsv"):
            df = pd.read_csv(path, sep="\t" if suffix == ".tsv" else ",", encoding="utf-8")
            return df.fillna("").to_string(index=False)
        if suffix in (".xlsx", ".xls"):
            sheets = []
            for sheet_name in pd.ExcelFile(path).sheet_names:
                df = pd.read_excel(path, sheet_name=sheet_name)
                sheets.append(f"### 工作表: {sheet_name} ###\n" + df.fillna("").to_string(index=False))
            return "\n\n".join(sheets)
        if suffix == ".json":
            return json.dumps(json.loads(path.read_text(encoding="utf-8")), ensure_ascii=False, indent=2)
        if suffix == ".jsonl":
            lines = [line for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]
            return "\n".join(json.dumps(json.loads(line), ensure_ascii=False) for line in lines)
        if suffix == ".docx":
            from docx import Document
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                parts.append(pd.DataFrame(rows).fillna("").to_string(index=False))
            return "\n".join(parts)
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        raise ToolException(f"文件 {path.name} 是二进制格式,无法直接读取,请转为 csv/txt 后使用")
    except Exception as e:
        raise ToolException(f"读取 {path.name} 失败: {e}")

@tool
def read_workspace(work_dir: str, rel_path: str = "") -> str:
    """
    读取工作区文件或目录内容。
    - rel_path 指向文件:返回文件内容(csv/xlsx/json/docx 等自动转换可读文本)。
    - rel_path 为空或指向目录:返回目录下的文件清单(名称+大小),适合查看 photo 等目录。
    参数:
        work_dir: 工作区目录,可选 code(解题代码)、paper(每道题思路与建模公式)、photo(模型生成的图片)。
        rel_path: 相对路径,如 "q1.py"、"问题1_思路.md";禁止绝对路径或 ".."。
    返回:
        文件内容或目录文件清单。
    """
    target = _resolve_path(work_dir, rel_path) if rel_path and rel_path.strip() else _resolve_path(work_dir, ".")
    if not rel_path or not rel_path.strip() or target.is_dir():
        if not target.is_dir():
            return f"({work_dir} 目录为空或不存在)"
        items = sorted(target.iterdir())
        if not items:
            return f"({work_dir} 目录为空)"
        lines = []
        for item in items:
            if item.is_file():
                lines.append(f"{item.name}  ({item.stat().st_size} 字节)")
            else:
                lines.append(f"{item.name}/  (文件夹)")
        return "\n".join(lines)
    if not target.is_file():
        raise ToolException(f"文件不存在: {target}")
    return f"=== {target.name} ===\n" + _read_smart(target)

@tool
def write_workspace(work_dir: str, rel_path: str, content: str) -> str:
    """
    新建或覆盖写入工作区文件(UTF-8 文本),父目录不存在时自动创建。
    参数:
        work_dir: code / paper / photo 之一。
        rel_path: 相对路径,如 "q1_solution.py"、"paper" 下建议 "问题1_思路.md"。
        content: 要写入的完整内容。
    返回:
        写入成功的完整路径。
    """
    target = _resolve_path(work_dir, rel_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content or "", encoding="utf-8")
    return f"已写入: {target}"

@tool
def append_workspace(work_dir: str, rel_path: str, content: str) -> str:
    """
    向工作区文件末尾追加内容(UTF-8),适合增量撰写论文思路或积累代码片段。
    参数:
        work_dir: code / paper / photo 之一。
        rel_path: 相对路径。
        content: 要追加的内容。
    返回:
        追加成功的信息。
    """
    target = _resolve_path(work_dir, rel_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    with open(target, "a", encoding="utf-8") as f:
        f.write(content or "")
    return f"已追加到: {target}"

@tool
def modify_workspace(work_dir: str, rel_path: str, old_str: str, new_str: str) -> str:
    """
    替换工作区文件中的一段文本(精确匹配一次),适合修改代码或修正论文表述。
    参数:
        work_dir: code / paper / photo 之一。
        rel_path: 相对路径。
        old_str: 要查找的原文,必须在文件中唯一出现。
        new_str: 替换后的新文本。
    返回:
        修改成功的信息。
    """
    target = _resolve_path(work_dir, rel_path)
    if not target.is_file():
        raise ToolException(f"文件不存在: {target}")
    if not old_str or not old_str.strip():
        raise ToolException(f"modify 必须提供 old_str")
    text = target.read_text(encoding="utf-8")
    count = text.count(old_str)
    if count == 0:
        raise ToolException(f"文件中找不到要替换的内容: {old_str[:50]}")
    if count > 1:
        raise ToolException(f"old_str 在文件中出现 {count} 次,请提供更长、更唯一的上下文片段")
    target.write_text(text.replace(old_str, new_str or "", 1), encoding="utf-8")
    return f"已修改: {target}"




tools=[search, read_workspace, write_workspace, append_workspace, modify_workspace]
